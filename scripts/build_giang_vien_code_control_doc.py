# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\xampp\htdocs\C-test--main\C-test-")
OUTPUT = ROOT / "Cau_hoi_giang_vien_code_control_VinhKhanhGuide.docx"
PRD = Path(r"C:\Users\Admin\Downloads\PRD_VinhKhanhGuide_WhiteMermaid.html")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_label_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label)
    r.bold = True
    r.font.color.rgb = RGBColor(154, 52, 18)
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.08
        p.add_run(item)


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for idx, line in enumerate(code.strip("\n").splitlines()):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.2)


def add_question(doc: Document, number: int, title: str, short: str, sequence: list[str], code: list[str], edit: list[str] | None = None, snippet: str | None = None) -> None:
    doc.add_heading(f"{number}. {title}", level=2)
    add_label_paragraph(doc, "Trả lời ngắn: ", short)
    add_label_paragraph(doc, "Nằm ở sequence: ", "")
    add_bullets(doc, sequence)
    add_label_paragraph(doc, "Method / file / dòng code: ", "")
    add_bullets(doc, code)
    if edit:
        add_label_paragraph(doc, "Nếu thầy bắt sửa live: ", "")
        add_bullets(doc, edit)
    if snippet:
        add_code(doc, snippet)


def build_doc() -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.4)
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(154, 52, 18)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12.2)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(10.5)
    styles["Heading 3"].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Bộ câu hỏi vấn đáp code-control\nVinhKhanhGuide")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(154, 52, 18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Bám PRD, sequence và code hiện tại. Dùng để trả lời khi giảng viên hỏi: nằm ở đâu trong sequence, method nào, code dòng nào, sửa ra sao.").italic = True

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = True
    for row in meta.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = [
        ("Code root", str(ROOT)),
        ("PRD", str(PRD)),
        ("Ngày soạn", "2026-04-29"),
        ("Cách dùng", "Mở đúng file theo path, chỉ vào line, đọc câu trả lời ngắn trước rồi nói phần sửa code nếu bị hỏi tiếp."),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell_shading(meta.cell(i, 0), "FFF7ED")
        set_cell_text(meta.cell(i, 0), k, bold=True)
        set_cell_text(meta.cell(i, 1), v)

    doc.add_paragraph()
    add_label_paragraph(doc, "Nguyên tắc trả lời với thầy: ", "không nói chung chung. Luôn đi theo chuỗi: sequence -> method -> code line -> behavior hiện tại -> nếu sửa thì sửa đúng điểm đó.")

    doc.add_heading("A. Bốn câu thầy đã từng hỏi", level=1)

    add_question(
        doc,
        1,
        "Thiết bị đang hoạt động bị hiện x2 thì viết lại code sao?",
        "Không phải lỗi view. Hiện tại API đang xem một phiên thiết bị là DeviceId + ClientInstanceId. App mở lại/resume tạo ClientInstanceId mới; nếu phiên cũ chưa disconnect hoặc chưa quá ActiveTimeout 30 giây thì Dashboard có thể đếm cùng một máy thành 2 session.",
        [
            "APP-12 Active device heartbeat: PRD lines 851-884. App StartAsync -> heartbeat -> API RegisterHeartbeat -> WebAdmin đọc count.",
            "WEB-02 Dashboard realtime: PRD lines 960-1000. Dashboard mở SSE / polling để lấy ActiveDeviceStatsDto.",
        ],
        [
            "VinhKhanhGuide.App/App.xaml.cs lines 55-73: OnStart/OnResume gọi StartAsync, OnSleep gọi StopAsync.",
            "VinhKhanhGuide.App/Services/ActiveDeviceTracker.cs lines 46-60: StartAsync tạo _clientInstanceId mới ở line 55.",
            "ActiveDeviceTracker.cs lines 108-145: SendHeartbeatAsync POST api/analytics/active-devices/heartbeat.",
            "ActiveDeviceTracker.cs lines 164-201: BuildHeartbeatRequestAsync gửi DeviceId + ClientInstanceId.",
            "VKFoodAPI/Controllers/ActiveDevicesController.cs lines 38-49: Heartbeat gọi _repository.RegisterHeartbeat.",
            "VKFoodAPI/Services/ActiveDeviceRepository.cs lines 83-132: RegisterHeartbeat upsert theo sessionKey.",
            "ActiveDeviceRepository.cs lines 189-217: BuildStatsUnsafe lấy activeDevices.Count để trả ActiveDeviceCount.",
            "ActiveDeviceRepository.cs lines 427-435: BuildSessionKey hiện đang ghép DeviceId:ClientInstanceId.",
            "CTest.WebAdmin/Views/Home/Dashboard.cshtml lines 298-303, 331-376, 617-633: view chỉ hiển thị count/list từ API.",
        ],
        [
            "Cách sửa nhanh nhất: sửa BuildSessionKey để 1 DeviceId chỉ có 1 session active. Khi app resume tạo client instance mới, heartbeat sẽ overwrite session cũ thay vì sinh bản ghi thứ hai.",
            "Nếu vẫn muốn giữ ClientInstanceId để debug, sửa BuildStatsUnsafe group by DeviceId rồi lấy bản LastSeenAtUtc mới nhất. View không cần sửa vì view đang nhận ActiveDeviceCount từ API.",
            "Sau khi sửa, test bằng cách mở app, refresh dashboard, sleep/resume trong dưới 30 giây: count vẫn là 1.",
        ],
        """// VKFoodAPI/Services/ActiveDeviceRepository.cs
private static string BuildSessionKey(string? deviceId, string? clientInstanceId)
{
    // 1 installation/device = 1 active device on dashboard.
    // ClientInstanceId can still be stored on the DTO for debugging,
    // but it should not split the same phone into two active sessions.
    return NormalizeDeviceId(deviceId);
}""",
    )

    add_question(
        doc,
        2,
        "Heatmap là gì? Map Analytics heatmap khác gì map bình thường? Xử lý heatmap trong code thế nào?",
        "Map bình thường hiển thị marker POI, bán kính POI và tuyến route. Heatmap là lớp phủ mật độ: vùng xanh/nhạt là ít điểm di chuyển, vàng/cam/đỏ là nhiều điểm thiết bị tụ lại. Trong code hiện tại heatmap được tính từ route points của active devices trong 12 giờ gần nhất; nếu chưa có route thì dùng vị trí online hiện tại. Nó không phải ảnh tĩnh, mà là overlay Leaflet heat.",
        [
            "WEB-05 Map Analytics: PRD lines 1108-1147. Service load POI/audio/history/active devices rồi view render POI map, heatmap, route.",
        ],
        [
            "CTest.WebAdmin/Services/PoiAdminService.cs lines 83-142: LoadMapManagementPageAsync load POI/audio/history/active devices.",
            "PoiAdminService.cs lines 305-365: ApplyMapAnalytics gán ActiveDeviceStats, Routes, HeatmapPoints, TopListeningPois.",
            "PoiAdminService.cs lines 423-469: BuildHeatmapPoints group tọa độ, tính Count và Weight.",
            "CTest.WebAdmin/Models/AdminViewModels.cs lines 88-109: MapPoiManagementViewModel có HeatmapPoints.",
            "AdminViewModels.cs lines 131-137: MapHeatPointViewModel gồm Latitude, Longitude, Count, Weight.",
            "CTest.WebAdmin/Views/MapPois/Index.cshtml lines 29-47: serialize POI marker data.",
            "MapPois/Index.cshtml lines 66-73: serialize heatmapJson.",
            "MapPois/Index.cshtml lines 253-255: import Leaflet và leaflet-heat.",
            "MapPois/Index.cshtml lines 294-325: vẽ marker POI và circle bán kính.",
            "MapPois/Index.cshtml lines 363-383: tạo L.heatLayer với radius, blur, gradient xanh -> đỏ.",
            "MapPois/Index.cshtml lines 394-403: add layer control POI / Tuyến thiết bị / Heatmap.",
        ],
        [
            "Muốn heatmap rõ hơn: chỉnh radius/blur/gradient ở MapPois/Index.cshtml lines 371-383.",
            "Muốn heatmap theo lượt nghe thay vì tuyến di chuyển: thêm hàm BuildListeningHeatmapPoints dùng history + tọa độ POI, rồi ở ApplyMapAnalytics line 327 đổi nguồn HeatmapPoints.",
            "Khi thầy hỏi nhận biết: chỉ vào layer control có checkbox Heatmap và vùng màu đỏ/cam trên bản đồ; marker POI vẫn là số tròn, route là đường polyline.",
        ],
        """// CTest.WebAdmin/Services/PoiAdminService.cs
// Đổi heatmap từ mật độ di chuyển sang mật độ lượt nghe theo POI:
vm.HeatmapPoints = BuildListeningHeatmapPoints(orderedHistory, vm.Pois, maxHeatPoints: 300);

private static List<MapHeatPointViewModel> BuildListeningHeatmapPoints(
    IReadOnlyList<ListeningHistoryEntryDto> history,
    IReadOnlyList<PoiListItemViewModel> pois,
    int maxHeatPoints)
{
    var poiById = pois.Where(p => p.Id != Guid.Empty)
        .ToDictionary(p => p.Id);

    var buckets = history
        .Where(h => poiById.ContainsKey(h.PoiId))
        .GroupBy(h => h.PoiId)
        .Select(g =>
        {
            var poi = poiById[g.Key];
            return new { poi.Latitude, poi.Longitude, Count = g.Count() };
        })
        .OrderByDescending(x => x.Count)
        .Take(Math.Clamp(maxHeatPoints, 50, 500))
        .ToList();

    var maxCount = buckets.Count == 0 ? 1 : buckets.Max(x => x.Count);
    return buckets.Select(x => new MapHeatPointViewModel
    {
        Latitude = x.Latitude,
        Longitude = x.Longitude,
        Count = x.Count,
        Weight = (double)x.Count / maxCount
    }).ToList();
}""",
    )

    add_question(
        doc,
        3,
        "Nếu thiết bị nằm giữa hai POI thì phát POI nào?",
        "Nếu đang đi tour thì app chỉ xét current stop của tour, nên current stop thắng dù POI khác gần hơn. Nếu không có tour, code tạo candidate trong bán kính rồi chọn theo Priority cao nhất, sau đó Distance gần nhất, sau đó Name, cuối cùng Id. Có thêm SamePrioritySwitchThresholdMeters = 8m để tránh GPS rung làm nhảy qua lại giữa hai POI cùng priority.",
        [
            "APP-03 GPS và geofence: PRD lines 463-510, đặc biệt lines 505-509 nói nhiều POI chồng vùng thì chọn best matched POI.",
            "APP-04 Auto narration: PRD lines 513-564, candidate hợp lệ mới PlayPoiNarrationAsync.",
            "APP-07 tour auto-advance: PRD lines 639-670, nếu đang tour thì dùng activeTourStopIndex/current stop.",
        ],
        [
            "VinhKhanhGuide.App/ViewModels/MainViewModel.cs lines 2519-2534: OnLocationUpdated khóa bằng _locationUpdateGate để tránh xử lý chồng.",
            "MainViewModel.cs lines 2552-2637: ApplyLocationSnapshotAsync gọi _geofenceEngine.Evaluate, refresh UI, resolve decision, rồi NarratePoiAsync.",
            "VinhKhanhGuide.App/Services/GeofenceEngine.cs lines 7-20: tính distance bằng GeoMath và IsInside = distance <= TriggerRadiusMeters.",
            "MainViewModel.cs lines 2987-3010: CreateAutoNarrationCandidates; nếu HasActiveTour thì filter chỉ current tour POI ở lines 2992-3001.",
            "VinhKhanhGuide.Core/Services/PoiAutoNarrationDecisionService.cs lines 8-21: CreateCandidates sort Priority desc -> Distance asc -> Name -> Id.",
            "PoiAutoNarrationDecisionService.cs lines 103-130: SelectStableCandidate giữ POI hiện tại nếu cùng priority và chênh lệch khoảng cách chưa vượt 8m.",
            "VinhKhanhGuide.Core/Models/POI.cs lines 33-37: Latitude, Longitude, TriggerRadiusMeters, Priority, CooldownMinutes.",
            "VinhKhanhGuide.Core.Tests/PoiAutoNarrationDecisionServiceTests.cs lines 52-86: test priority cao thắng và cùng priority thì gần hơn thắng.",
        ],
        [
            "Nếu muốn nearest luôn thắng dù priority thấp hơn: đổi sort ở CreateCandidates và NormalizeCandidates thành Distance asc trước, Priority desc sau.",
            "Nếu muốn giảm nhảy POI khi đứng giữa hai quán: tăng SamePrioritySwitchThresholdMeters ở MainViewModel.cs line 29.",
            "Nếu muốn tour không ép current stop: bỏ filter HasActiveTour ở MainViewModel.cs lines 2992-3001.",
        ],
        """// VinhKhanhGuide.Core/Services/PoiAutoNarrationDecisionService.cs
// Sửa nếu muốn POI gần nhất thắng trước priority:
.OrderBy(candidate => candidate.DistanceMeters)
.ThenByDescending(candidate => candidate.Poi.Priority)
.ThenBy(candidate => candidate.Poi.Name, StringComparer.OrdinalIgnoreCase)
.ThenBy(candidate => candidate.Poi.Id)""",
    )

    add_question(
        doc,
        4,
        "Nếu nhiều thiết bị cùng truy cập một POI thì sắp xếp hàng đợi như thế nào?",
        "Có hai lớp cần phân biệt. Âm thanh phát local trên từng thiết bị nên không có server queue để chặn người sau nghe sau. Nhưng server có TtsQueuePosition trong ListeningHistory để đánh số thứ tự phiên nghe theo từng POI. Thứ tự hiện tại là StartedAtUtc sớm hơn trước, nếu trùng thì ReceivedAtUtc, nếu vẫn trùng thì Id.",
        [
            "APP-04 lines 551-556: sau khi phát tạo/update listening session.",
            "APP-10 Listening History: PRD lines 768-809.",
            "WEB-07 QR public: PRD lines 1227-1233 cũng POST/PUT listening-history khi khách nghe qua QR web.",
        ],
        [
            "VinhKhanhGuide.App/ViewModels/MainViewModel.cs lines 3388-3407: NarratePoiAsync ghi optimistic history và gọi BeginAsync.",
            "VinhKhanhGuide.App/Services/ListeningHistorySyncService.cs lines 26-77: BeginAsync POST api/analytics/listening-history.",
            "VKFoodAPI/Controllers/ListeningHistoryController.cs lines 39-44: Create gọi repository.Create.",
            "VKFoodAPI/Services/ListeningHistoryRepository.cs lines 80-90: Create add item rồi RecalculateQueuePositionsUnsafe(created.PoiId).",
            "ListeningHistoryRepository.cs lines 448-465: queue order = StartedAtUtc -> ReceivedAtUtc -> Id, gán TtsQueuePosition = index + 1.",
            "VinhKhanhGuide.Core/Contracts/ListeningHistoryContracts.cs line 27: DTO có TtsQueuePosition.",
            "CTest.WebAdmin/Services/ListeningHistoryService.cs lines 161-186: map DTO sang TimelineItem giữ TtsQueuePosition.",
            "CTest.WebAdmin/Views/UsageLogs/_UsageHistoryContent.cshtml lines 118-124 và 159-162: hiển thị cột TTS queue.",
        ],
        [
            "Nếu thầy muốn queue chỉ tính các phiên đang nghe, sửa RecalculateQueuePositionsUnsafe filter thêm !item.Completed.",
            "Sau khi sửa queue active-only, Update phải gọi RecalculateQueuePositionsUnsafe(existing.PoiId) sau khi Completed thay đổi, nếu không người đã nghe xong vẫn còn chiếm số.",
            "Nếu muốn queue thật sự điều phối phát âm thanh, phải thêm endpoint acquire/release queue trước khi app gọi NarrateAsync; hiện code chưa làm vì mỗi thiết bị phát độc lập.",
        ],
        """// VKFoodAPI/Services/ListeningHistoryRepository.cs
private void RecalculateQueuePositionsUnsafe(Guid poiId)
{
    var ordered = _items
        .Where(item => item.PoiId == poiId && !item.Completed)
        .OrderBy(item => item.StartedAtUtc)
        .ThenBy(item => item.ReceivedAtUtc)
        .ThenBy(item => item.Id)
        .ToList();

    for (var index = 0; index < ordered.Count; index++)
        ordered[index].TtsQueuePosition = index + 1;
}

// Trong Update(), sau khi set Completed:
var affectedPoiId = existing.PoiId;
RecalculateQueuePositionsUnsafe(affectedPoiId);
SaveUnsafe();""",
    )

    doc.add_heading("B. Câu hỏi bổ sung giảng viên có thể hỏi", level=1)

    extra_questions = [
        (
            "GPS khác QR như thế nào?",
            "GPS cần quyền vị trí và xét bán kính POI; QR không cần GPS, chỉ resolve mã để mở POI/Tour. Nếu chưa login, app có thể vào guest để mở QR không bị gián đoạn.",
            ["APP-08 QR: PRD lines 681-726.", "WEB-07 QR public: PRD lines 1196-1236."],
            [
                "VinhKhanhGuide.App/Views/QrScannerPage.xaml.cs lines 85-98: nhận barcode, chống xử lý trùng bằng _isHandlingResult.",
                "QrScannerPage.xaml.cs lines 100-134: Resolve QR, nếu Tour thì OpenTourFromQrAsync, nếu POI thì OpenPoiFromQrAsync.",
                "VinhKhanhGuide.App/Services/QrResolveService.cs lines 21-47: gọi GET api/resolve-qr.",
                "MainViewModel.cs lines 1893-1981: OpenPoiFromQrAsync mở POI, không yêu cầu GPS; lines 1959-1966 set LocationText 'Không cần bật vị trí'.",
                "VinhKhanhGuide.App/App.xaml.cs lines 136-176 và 191-223: deep link QR tự guest login rồi mở target.",
            ],
            ["Muốn QR chỉ mở chi tiết, không tự phát: ở QrScannerPage.xaml.cs line 127 gọi OpenPoiFromQrAsync(poiId, autoPlay: false)."],
            """// VinhKhanhGuide.App/Views/QrScannerPage.xaml.cs
if (!Guid.TryParse(resolved.TargetId, out var poiId) ||
    !await _viewModel.OpenPoiFromQrAsync(poiId, autoPlay: false))
{
    await ShowScanErrorAsync("Không mở được POI từ mã QR này.");
    return;
}""",
        ),
        (
            "Auto play khác manual play ở đâu?",
            "Auto play do GPS/geofence kích hoạt, có debounce/cooldown và không chen ngang khi đang phát. Manual play do người dùng bấm nghe, đi thẳng vào NarratePoiAsync với autoTriggered=false.",
            ["APP-04 Auto narration: PRD lines 513-564.", "APP-05 Manual POI playback: PRD lines 565-603."],
            [
                "MainViewModel.cs lines 2603-2616: auto narration chỉ chạy nếu allowAutoNarrate và IsAutoNarrationEnabled.",
                "PoiAutoNarrationDecisionService.cs lines 24-78: Decide kiểm tra debounce, candidate, đang phát, cooldown.",
                "MainViewModel.cs lines 2065-2085: ToggleSelectedPoiNarrationAsync cho manual selected POI.",
                "MainViewModel.cs lines 2087-2106: TogglePoiNarrationAsync cho manual từ list/map.",
                "MainViewModel.cs lines 3328-3485: NarratePoiAsync dùng chung cho cả auto/manual; autoTriggered quyết định trigger history GPS hay APP.",
                "NarrationService.cs lines 54-64: phát mới cancel speech cũ.",
            ],
            ["Muốn manual luôn dừng auto trước khi phát: thêm StopNarrationAsync trước NarratePoiAsync trong TogglePoiNarrationAsync nếu đang có narration khác."],
            """// VinhKhanhGuide.App/ViewModels/MainViewModel.cs
public async Task TogglePoiNarrationAsync(Guid poiId)
{
    var poi = _pois.FirstOrDefault(item => item.Id == poiId);
    if (poi is null) return;

    if (IsCurrentNarration(poi))
    {
        await StopNarrationAsync();
        return;
    }

    if (IsNarrating)
    {
        await StopNarrationAsync();
    }

    await NarratePoiAsync(poi, false, GetDistanceForPoi(poi.Id), syncSelectedPoi: false);
}""",
        ),
        (
            "Tour làm sao tránh phát sai POI gần hơn?",
            "Tour dùng activeTourStopIndex. Khi đang tour, candidate auto narration chỉ còn current stop, nên POI ngoài current stop không được tự phát dù gần hơn.",
            ["APP-06 chọn tour: PRD lines 604-638.", "APP-07 đi theo tour: PRD lines 639-670."],
            [
                "MainViewModel.cs lines 1668-1728: ActivateTourAsync set _activeTourId và _activeTourStopIndex = 0.",
                "MainViewModel.cs lines 2878-2898: GetTourStopIds/GetActiveTourStopIds lấy danh sách stop hợp lệ.",
                "MainViewModel.cs lines 2900-2914: GetCurrentActiveTourPoiId/GetCurrentActiveTourPoi trả current stop.",
                "MainViewModel.cs lines 2987-3010: CreateAutoNarrationCandidates filter item.Poi.Id == currentTourPoiId khi HasActiveTour.",
                "MainViewModel.cs lines 2943-2975: TryAdvanceActiveTourAfterNarration tăng _activeTourStopIndex sau khi nghe xong.",
                "MainViewModel.cs lines 3449-3468: NarratePoiAsync gọi TryAdvanceActiveTourAfterNarration.",
            ],
            ["Muốn chỉ auto mới làm tour advance, manual nghe thử không làm nhảy chặng: đổi ShouldAdvanceActiveTourAfterNarration chỉ return autoTriggered."],
            """// VinhKhanhGuide.App/ViewModels/MainViewModel.cs
private static bool ShouldAdvanceActiveTourAfterNarration(
    POI poi,
    bool autoTriggered,
    double? distanceMeters)
{
    return autoTriggered;
}""",
        ),
        (
            "Nếu API lỗi thì app có vỡ demo không?",
            "App vẫn có fallback: API thành công thì cache snapshot; API lỗi thì đọc SQLite offline snapshot; nếu chưa có thì dùng last successful snapshot; cuối cùng dùng seed data.",
            ["APP-01/APP-09 trong PRD mô tả app tải POI/tour và fallback khi xem detail."],
            [
                "VinhKhanhGuide.App/Services/PoiProvider.cs lines 40-92: GetPoisAsync API -> offline snapshot -> cached snapshot -> seed.",
                "PoiProvider.cs lines 94-137: GetPoiByIdAsync API -> offline -> cached -> seed.",
                "PoiProvider.cs lines 148-160: PersistOfflineSnapshotAsync lưu snapshot vào offline store.",
                "VinhKhanhGuide.App/Services/PoiOfflineStore.cs lines 29-56: đọc POI SQLite.",
                "PoiOfflineStore.cs lines 88-130: ReplacePoisAsync ghi snapshot.",
                "VinhKhanhGuide.App/Services/TourProvider.cs lines 24-53: tour API -> cached -> seed.",
            ],
            ["Muốn demo luôn báo nguồn dữ liệu đang dùng: hiển thị PoiRepository.CurrentDataSource trên UI sau khi InitializeAsync load POI."],
            None,
        ),
        (
            "WebAdmin phân quyền Admin và chủ quán ở đâu?",
            "Web dùng Cookie Auth. Admin vào Dashboard/POI/Map; PoiOwner vào Owner Portal và chỉ thấy POI của mình thông qua CurrentUser filter.",
            ["WEB-01 login/role: PRD lines 929-959.", "WEB-09 Owner Portal: PRD lines 1279-1316."],
            [
                "CTest.WebAdmin/Program.cs lines 18-27: cấu hình cookie CTest.WebAdmin.Auth.",
                "Program.cs lines 32-43: policy AdminOnly và OwnerArea.",
                "CTest.WebAdmin/Security/WebAdminSecurity.cs lines 8-18: WebAdminRoles và WebAdminPolicies.",
                "AccountController.cs lines 38-73: validate login, tạo claims Role/OwnerCode/OwnerEmail.",
                "AccountController.cs lines 91-104: Admin redirect Home, owner redirect Owner.",
                "PoiAdminService.cs lines 524-539: FilterPoisForCurrentUser chỉ trả POI owner được quản lý.",
                "MapPoisController.cs line 9 và PoisController.cs line 10: hiện đang AdminOnly.",
            ],
            ["Nếu muốn owner được mở Map Analytics cho POI của họ: đổi MapPoisController attribute từ AdminOnly sang OwnerArea; service đã có FilterPoisForCurrentUser."],
            """// CTest.WebAdmin/Controllers/MapPoisController.cs
[Authorize(Policy = WebAdminPolicies.OwnerArea)]
public class MapPoisController : Controller
{
    // PoiAdminService.FilterPoisForCurrentUser() đã lọc POI theo owner.
}""",
        ),
        (
            "Dashboard lấy số liệu từ đâu, có tự refresh không?",
            "Dashboard gọi nhiều API song song để build lần đầu, sau đó active devices cập nhật bằng SSE 1 giây hoặc polling 3 giây; usage history snapshot cũng polling.",
            ["WEB-02 Dashboard: PRD lines 960-1006."],
            [
                "DashboardService.cs lines 31-51: LoadAsync gọi POI, tour, audio, history, active devices song song.",
                "DashboardService.cs lines 145-230: BuildFromSharedData tính TotalPois, TodayListenCount, TopPois, ActiveDeviceCount.",
                "HomeController.cs lines 46-80: ActiveDeviceEvents SSE mỗi 1 giây, chỉ gửi khi payload đổi.",
                "Dashboard.cshtml lines 617-633: update active device count/list trên UI.",
                "Dashboard.cshtml lines 672-702: start polling 3 giây và EventSource SSE.",
            ],
            ["Nếu muốn giảm tải API: tăng interval ở Dashboard.cshtml line 682 từ 3000 lên 8000 hoặc trong HomeController line 53 tăng PeriodicTimer."],
            None,
        ),
        (
            "Map app vẽ marker POI như thế nào?",
            "MainPage dùng Mapsui MapView. Mỗi lần state đổi, RefreshMapPins xóa pin cũ, thêm entrance, user pin, route tour, rồi thêm POI pin theo PreviewMapPoiStatuses/VisibleMapPoiStatuses.",
            ["APP-03 và APP-07: map cập nhật marker theo GPS/tour."],
            [
                "MainPage.xaml.cs lines 82-93: InitializeMapsui tạo map và gắn event pin/click.",
                "MainPage.xaml.cs lines 484-486: OnMapStateChanged gọi RefreshMapPins.",
                "MainPage.xaml.cs lines 497-568: RefreshMapPins xóa pin, add entrance/user/route/POI, center map.",
                "MainPage.xaml.cs lines 570-590: RefreshTourRoute vẽ polyline tour.",
                "MainPage.xaml.cs lines 929-958: CreateRestaurantPin đổi màu theo selected/current tour/completed/inside/nearest.",
            ],
            ["Muốn đổi màu POI đang trong bán kính: sửa Color ở CreateRestaurantPin lines 935-945."],
            None,
        ),
        (
            "Muốn đổi thời gian heartbeat hoặc timeout active devices thì sửa ở đâu?",
            "App gửi heartbeat mỗi 8 giây, server coi active trong 30 giây. Hai con số này nằm ở hai project khác nhau.",
            ["APP-12 Active device heartbeat: PRD lines 851-884."],
            [
                "ActiveDeviceTracker.cs line 16: HeartbeatInterval = 8 giây.",
                "ActiveDeviceTracker.cs lines 147-156: RunHeartbeatLoopAsync delay theo HeartbeatInterval.",
                "ActiveDeviceRepository.cs line 11: ActiveTimeout = 30 giây.",
                "ActiveDeviceRepository.cs lines 219-233: PruneInactiveUnsafe xóa session quá timeout.",
                "ActiveDeviceRepository.cs lines 189-217: BuildStatsUnsafe chỉ tính LastSeenAtUtc >= threshold.",
            ],
            ["Nếu tăng heartbeat lên 15 giây thì nên tăng timeout ít nhất 45-60 giây để tránh dashboard nhấp nháy offline khi mạng yếu."],
            """// App
private static readonly TimeSpan HeartbeatInterval = TimeSpan.FromSeconds(15);

// API
private static readonly TimeSpan ActiveTimeout = TimeSpan.FromSeconds(60);""",
        ),
        (
            "Audio file và TTS được chọn như thế nào?",
            "NarrationService ưu tiên audio nếu playbackMode=audio và có audioAssetPath; nếu không thì dùng TextToSpeech với locale theo language. MainViewModel tạo playback request trước khi gọi NarrationService.",
            ["APP-04 lines 543-548: có audio thì play file/cache, không có thì SpeakTextAsync.", "APP-11 settings: PRD lines 810-849."],
            [
                "MainViewModel.cs lines 3349-3352: ResolvePlaybackRequest chọn mode/audio path.",
                "MainViewModel.cs lines 3411-3415: gọi _narrationService.NarrateAsync(poi, language, playbackMode).",
                "NarrationService.cs lines 25-39: NarrateAsync lấy text và audio path theo language.",
                "NarrationService.cs lines 41-77: SpeakAsync, mode audio thì PlayAudioAsync, còn lại TextToSpeech.Default.SpeakAsync.",
                "NarrationService.cs lines 112-170: chọn locale theo vi/en/zh/ko/fr.",
            ],
            ["Muốn luôn fallback TTS khi file audio lỗi: catch lỗi PlayAudioAsync trong SpeakAsync rồi gọi TextToSpeech.Default.SpeakAsync(text, options)."],
            None,
        ),
        (
            "Lịch sử nghe và top POI tính như thế nào?",
            "Mỗi lần NarratePoiAsync bắt đầu thì tạo listening history, khi phát xong thì update completed/listen seconds. Dashboard và UsageLogs group theo PoiId/PoiName để tính top POI.",
            ["APP-10 Listening History: PRD lines 768-809.", "WEB-02 Dashboard: PRD lines 987-1004."],
            [
                "MainViewModel.cs lines 3388-3407: add local optimistic history và BeginAsync.",
                "MainViewModel.cs lines 3437-3447: CompleteAsync cập nhật listenSeconds/completed.",
                "ListeningHistorySyncService.cs lines 26-77: POST create.",
                "ListeningHistorySyncService.cs lines 79-115: PUT complete.",
                "ListeningHistoryRepository.cs lines 30-49: GetListeningHistory filter/sort/limit.",
                "ListeningHistoryRepository.cs lines 52-78: CountListeningByPoi group theo POI.",
                "DashboardService.cs lines 164-174: topPois lấy từ localizedHistory GroupBy.",
            ],
            ["Nếu muốn Top POI chỉ tính completed=true: thêm .Where(x => x.Item.Completed) trước GroupBy trong DashboardService.cs lines 164-174."],
            None,
        ),
        (
            "POI owner đăng ký POI mới thì vì sao chưa hiện app ngay?",
            "Owner submission được tạo với IsActive=false để chờ Admin duyệt. App chỉ lấy DTO IsActive=true, nên chưa duyệt thì không hiện trên app.",
            ["WEB-03 POI CRUD: PRD lines 1007-1062.", "WEB-09 Owner Portal: PRD lines 1309-1316."],
            [
                "PoiAdminService.cs lines 172-203: CreateAsync; nếu owner submission thì ApplyOwner và dto.IsActive=false ở lines 189-194.",
                "PoisController.cs lines 144-159: Approve action gọi ApproveAsync.",
                "PoiAdminService.cs lines 205-236: ApproveAsync chỉ Admin, set poi.IsActive=true.",
                "PoiProvider.cs lines 44-48 và 202-209: app chỉ map/filter POI IsActive.",
            ],
            ["Nếu muốn owner POI tự active không cần duyệt: bỏ dto.IsActive=false trong PoiAdminService.cs line 193, nhưng khi bảo vệ đồ án nên giữ duyệt để hợp lý nghiệp vụ."],
            None,
        ),
        (
            "Vì sao dùng JSON repository thay vì database?",
            "Phạm vi demo/đồ án dùng JSON trong App_Data để dễ chạy local, dễ inspect, không cần setup DB. Repository vẫn tách riêng nên sau này đổi sang database chỉ thay lớp repository/API, không đổi app/web nhiều.",
            ["PRD overview lines 43 và 124 nói backend ASP.NET Core/VKFoodAPI, JSON repositories trong App_Data."],
            [
                "VKFoodAPI/Program.cs lines 23-33: đăng ký các repository singleton và hosted service.",
                "ActiveDeviceRepository.cs lines 31-32: active-devices.json và active-device-routes.json.",
                "ListeningHistoryRepository.cs lines 23-27: listening-history.json.",
                "ActiveDeviceRepository.cs lines 299-307 và ListeningHistoryRepository.cs lines 192-196: SaveUnsafe serialize JSON.",
            ],
            ["Nếu đổi sang DB, giữ interface public method như GetStats/RegisterHeartbeat/Create/Update, thay implementation lưu EF Core; controller không cần đổi nhiều."],
            None,
        ),
        (
            "Có test chứng minh chọn POI/cooldown không?",
            "Có test unit cho PoiAutoNarrationDecisionService: không candidate, priority thắng, cùng priority nearest thắng, cooldown, không phát trùng khi đang phát, chống GPS jitter.",
            ["APP-03/APP-04 liên quan geofence và auto narration."],
            [
                "VinhKhanhGuide.Core.Tests/PoiAutoNarrationDecisionServiceTests.cs lines 20-34: no candidate.",
                "Tests lines 52-86: priority và nearest.",
                "Tests lines 114-132: cooldown.",
                "Tests lines 134-169: không chen narration đang phát/không duplicate.",
                "Tests lines 171-207: GPS jitter không oscillation.",
            ],
            ["Nếu sửa logic chọn POI, phải sửa hoặc thêm test ở file này trước khi demo lại."],
            None,
        ),
    ]

    for idx, data in enumerate(extra_questions, start=5):
        title, short, sequence, code, edit, snippet = data
        add_question(doc, idx, title, short, sequence, code, edit, snippet)

    doc.add_heading("C. Mẫu trả lời rất ngắn khi bị hỏi gấp", level=1)
    quick = [
        ("Active devices x2", "Do server key đang là DeviceId + ClientInstanceId. Sửa ở ActiveDeviceRepository.BuildSessionKey hoặc group count theo DeviceId trong BuildStatsUnsafe; view không phải nơi sửa."),
        ("Heatmap", "Heatmap là lớp mật độ route/device positions, vẽ bằng leaflet-heat. Normal map là marker POI + circle + route polyline."),
        ("Đứng giữa POI", "Tour thì current stop thắng; không tour thì Priority desc, Distance asc, Name, Id. Code ở CreateAutoNarrationCandidates và PoiAutoNarrationDecisionService."),
        ("Nhiều thiết bị vào POI", "Không có queue phát âm thanh chung. Có TtsQueuePosition cho lịch sử nghe theo POI, sắp bằng StartedAtUtc, ReceivedAtUtc, Id."),
        ("QR", "QR không cần GPS. QrScannerPage resolve code rồi MainViewModel mở POI/Tour; deep link có thể tự guest login."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_shading(table.cell(0, 0), "9A3412")
    set_cell_shading(table.cell(0, 1), "9A3412")
    set_cell_text(table.cell(0, 0), "Chủ đề", True)
    set_cell_text(table.cell(0, 1), "Câu trả lời 1-2 câu", True)
    for k, v in quick:
        cells = table.add_row().cells
        set_cell_text(cells[0], k, True)
        set_cell_text(cells[1], v)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
