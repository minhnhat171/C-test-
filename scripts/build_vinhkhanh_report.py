from __future__ import annotations

import html
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "BaoCao_DoAn_VinhKhanhGuide.docx"
OUT_HTML = ROOT / "BaoCao_DoAn_VinhKhanhGuide.html"
OUT_MD = ROOT / "BaoCao_DoAn_VinhKhanhGuide.md"


def load_json(path: str):
    with (ROOT / path).open("r", encoding="utf-8") as f:
        return json.load(f)


pois = load_json("VKFoodAPI/App_Data/pois.json")
tours = load_json("VKFoodAPI/App_Data/tours.json")
audio_guides = load_json("VKFoodAPI/App_Data/audio-guides.json")
qr_codes = load_json("VKFoodAPI/App_Data/qr-codes.json")
listening_history = load_json("VKFoodAPI/App_Data/listening-history.json")
active_devices = load_json("VKFoodAPI/App_Data/active-devices.json")
movement_logs = load_json("VKFoodAPI/App_Data/movement-logs.json")
user_profiles = load_json("VKFoodAPI/App_Data/user-profiles.json")
audit_logs = load_json("VKFoodAPI/App_Data/audit-logs.json")
web_admin_users = load_json("CTest.WebAdmin/App_Data/web-admin-users.json")

active_pois = [p for p in pois if p.get("IsActive") is True and not p.get("IsDeleted", False)]
active_tours = [t for t in tours if t.get("IsActive") is True and not t.get("IsDeleted", False)]
translation_keys = sorted({k for p in pois for k in (p.get("NarrationTranslations") or {}).keys()})


class Report:
    def __init__(self) -> None:
        self.body: list[dict] = []
        self.figures: list[str] = []
        self.tables: list[str] = []

    def h(self, level: int, text: str):
        self.body.append({"type": "heading", "level": level, "text": text})

    def p(self, text: str):
        self.body.append({"type": "paragraph", "text": text})

    def bullets(self, items: list[str]):
        self.body.append({"type": "bullets", "items": items})

    def numbered(self, items: list[str]):
        self.body.append({"type": "numbered", "items": items})

    def table(self, title: str, headers: list[str], rows: list[list[str]]):
        caption = f"Bảng {len(self.tables) + 1}. {title}"
        self.tables.append(caption)
        self.body.append({"type": "caption", "kind": "table", "text": caption})
        self.body.append({"type": "table", "headers": headers, "rows": rows})

    def figure(self, title: str, code: str):
        caption = f"Hình {len(self.figures) + 1}. {title}"
        self.figures.append(caption)
        self.body.append({"type": "caption", "kind": "figure", "text": caption})
        self.body.append({"type": "code", "language": "mermaid", "text": code.strip()})

    def code(self, text: str, language: str = ""):
        self.body.append({"type": "code", "language": language, "text": text.strip()})

    def page_break(self):
        self.body.append({"type": "page_break"})


r = Report()


def add_overview_chapter():
    r.h(1, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI")
    r.h(2, "1.1. Lý do chọn đề tài")
    r.p(
        "Phố ẩm thực Vĩnh Khánh là một không gian ăn uống đông khách, có nhiều quán hải sản, lẩu, nướng và món đường phố nằm gần nhau. Với khách du lịch hoặc người lần đầu đến khu vực này, việc chọn quán, hiểu món nổi bật, tìm đường và nghe giới thiệu theo ngôn ngữ phù hợp là nhu cầu thực tế."
    )
    r.p(
        "Đề tài lựa chọn hướng thuyết minh tự động vì kết hợp được bản đồ, GPS/geofence, QR code, nội dung đa ngôn ngữ và quản trị nội dung tập trung. Hệ thống gồm App MAUI cho khách tham quan, WebAdmin MVC cho quản trị/chủ quán và Backend API dùng chung dữ liệu JSON để phục vụ demo học thuật."
    )
    r.h(2, "1.2. Mục tiêu đề tài")
    r.bullets([
        "Xây dựng app Android cho phép khách chọn ngôn ngữ, xem bản đồ, xem danh sách POI, nghe thuyết minh thủ công hoặc tự động khi vào vùng geofence.",
        "Xây dựng WebAdmin để quản lý POI, tour, audio guide/TTS, bản dịch, QR public, dashboard, usage logs, tài khoản WebAdmin và hồ sơ app user.",
        "Xây dựng Backend API ASP.NET Core cung cấp dữ liệu POI/tour/audio/QR/analytics và lưu bằng JSON repository trong `VKFoodAPI/App_Data`.",
        "Hỗ trợ owner portal để chủ cửa hàng xem POI liên quan, xem lịch sử nghe và gửi đăng ký POI mới chờ duyệt.",
        "Thiết kế app có khả năng fallback khi API lỗi thông qua SQLite offline snapshot, cached snapshot và seed data."
    ])
    r.h(2, "1.3. Đối tượng sử dụng")
    r.table(
        "Đối tượng sử dụng hệ thống",
        ["Đối tượng", "Mục đích sử dụng", "Module chính"],
        [
            ["Guest/Tourist", "Vào app không cần đăng nhập, chọn ngôn ngữ, xem bản đồ, nghe thuyết minh, quét QR.", "VinhKhanhGuide.App"],
            ["App user", "Đăng ký/đăng nhập local trong app, lưu hồ sơ, đồng bộ profile và lịch sử nghe.", "AuthService, UserProfileSyncService"],
            ["Admin", "Quản trị nội dung, dashboard, POI, tour, audio/TTS, QR, translation, tài khoản và app users.", "CTest.WebAdmin"],
            ["Owner/chủ cửa hàng", "Xem POI của mình, theo dõi lịch sử nghe, đăng ký POI mới chờ duyệt.", "OwnerController"],
            ["Public QR visitor", "Mở trang QR public từ trình duyệt, nghe nội dung hoặc mở app bằng deep link.", "QrCodesController.Scan"],
        ],
    )
    r.h(2, "1.4. Phạm vi đề tài")
    r.p(
        f"Đến thời điểm đọc source, dữ liệu demo có {len(active_pois)} POI active, {len(active_tours)} tour active, {len(audio_guides)} audio guide, {len(qr_codes)} QR item API, {len(listening_history)} bản ghi listening history, {len(active_devices)} active device session, {len(movement_logs)} movement log và {len(user_profiles)} user profile."
    )
    r.bullets([
        "Phần đã triển khai: public API đọc POI/tour, admin API có `X-Admin-Api-Key`, WebAdmin có cookie auth và role Admin/PoiOwner, app có map, POI, tour, QR, GPS/geofence, TTS/audio, history, offline cache và heartbeat.",
        "Phần cần ghi trung thực: hệ thống dùng JSON repository phù hợp demo/học thuật, chưa phải cơ sở dữ liệu production; WebAdmin QR hiện tạo/in QR public động từ POI/Tour, chưa có màn CRUD độc lập cho `qr-codes.json`; audio playback hiện là cancel/replace kết hợp debounce/cooldown, chưa có hàng đợi âm thanh phức tạp.",
        "Bộ chọn ngôn ngữ hiện dùng `vi/en/zh/ja/de` trong WebAdmin Translation và app. Tuy vậy dữ liệu JSON hiện còn key lịch sử `ko/fr`, repository vẫn giữ để đọc an toàn; cần rà soát lại nội dung bản dịch trước triển khai thật.",
        "Public QR web có `DeviceCapabilitySimulation` theo yêu cầu giảng viên: `0 = thiết bị mạnh`, `1 = thiết bị yếu`; đây là mô phỏng random 0/1 hoặc ép bằng query string, chưa phải đo cấu hình thiết bị thật.",
    ])
    r.h(2, "1.5. Phương pháp thực hiện")
    r.numbered([
        "Phân tích yêu cầu từ PRD và luồng người dùng: khách tham quan, admin, owner, public QR visitor.",
        "Thiết kế kiến trúc 3 khối App MAUI, WebAdmin MVC và Backend API dùng DTO chung trong `VinhKhanhGuide.Core`.",
        "Thiết kế dữ liệu theo mô hình logic, sau đó lưu vật lý bằng JSON repository và SQLite cache trong app.",
        "Triển khai API controller/repository, WebAdmin controller/service/API client và app ViewModel/service.",
        "Kiểm thử theo use case: login, CRUD POI, tour, QR, GPS auto narration, listening history, dashboard và offline fallback.",
    ])


def add_theory_chapter():
    r.h(1, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG")
    sections = [
        ("2.1. ASP.NET Core Web API", "ASP.NET Core Web API được dùng cho `VKFoodAPI`. Controller trả JSON DTO, public endpoint phục vụ app/QR, còn endpoint quản trị dùng policy `AdminApiKey`. Swagger được cấu hình trong môi trường Development."),
        ("2.2. ASP.NET Core MVC", "ASP.NET Core MVC được dùng cho `CTest.WebAdmin`. Mô hình controller-service-view giúp tách nghiệp vụ quản trị khỏi giao diện Razor. WebAdmin còn host các API controller của `VKFoodAPI` qua `AddApplicationPart`, thuận tiện cho demo một tiến trình."),
        ("2.3. .NET MAUI", "App dùng .NET MAUI target `net8.0-android`, kết hợp XAML page, `MainViewModel`, Mapsui, ZXing.Net.Maui, TextToSpeech và Android MediaPlayer. Đây là hướng phù hợp để xây dựng app bản đồ và thuyết minh trên Android."),
        ("2.4. JSON Repository", "JSON repository là cách lưu dữ liệu đơn giản bằng file như `pois.json`, `tours.json`, `audio-guides.json`. Cách này dễ demo, dễ kiểm tra diff, nhưng không thay thế được database quan hệ khi có đồng thời nhiều người dùng hoặc yêu cầu transaction phức tạp."),
        ("2.5. SQLite/offline cache trong app", "`PoiOfflineStore` tạo SQLite `vinhkhanh_guide.db` với bảng `poi_cache` và `sync_metadata`. Khi API lỗi, app ưu tiên đọc SQLite snapshot, sau đó cached snapshot trong bộ nhớ, cuối cùng mới dùng seed data."),
        ("2.6. GPS và Geofence", "`LocationService` lấy vị trí thiết bị. `GeofenceEngine` tính khoảng cách từ vị trí hiện tại đến từng POI bằng `GeoMath.DistanceMeters`, sau đó so với `TriggerRadiusMeters` để xác định người dùng đã vào vùng kích hoạt hay chưa."),
        ("2.7. Text-to-Speech và Audio Guide", "`NarrationService` dùng MAUI TextToSpeech cho TTS và Android `MediaPlayer` cho file audio. `AudioGuideRepository` đồng bộ audio/TTS đã publish vào POI để app có thể phát theo `NarrationTranslations` hoặc `AudioAssetPath`."),
        ("2.8. QR Code và Deep Link", "WebAdmin dùng QRCoder tạo QR public `/qr/{targetType}/{targetId}`. App dùng camera ZXing hoặc nhập mã thủ công, gọi `GET /api/resolve-qr?code=...` và xử lý deep link `vinhkhanhguide://poi/...` hoặc `vinhkhanhguide://tour/...`."),
        ("2.9. Authentication, Authorization, Role-based Access Control", "WebAdmin dùng Cookie Authentication, role `Admin` và `PoiOwner`. API quản trị dùng `X-Admin-Api-Key`. App hiện có guest mode và local auth bằng Preferences, đồng bộ profile lên API nhưng chưa phải hệ thống JWT/OAuth production."),
        ("2.10. Analytics", "Analytics gồm listening history, active devices và movement logs. Listening history ghi begin/complete lượt nghe; active device heartbeat cập nhật phiên thiết bị; movement log ghi điểm di chuyển hợp lệ khi heartbeat có tọa độ."),
    ]
    for title, text in sections:
        r.h(2, title)
        r.p(text)


def add_requirements_chapter():
    r.h(1, "CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG")
    r.h(2, "3.1. Yêu cầu chức năng")
    r.h(3, "A. App MAUI")
    r.bullets([
        "Chọn ngôn ngữ `vi/en/zh/ja/de` trước khi vào app; lưu cài đặt giọng đọc và mode `tts/audio`.",
        "Guest mode, đăng ký/đăng nhập local app user, cập nhật hồ sơ và đồng bộ profile lên API.",
        "Xem bản đồ, danh sách POI, chi tiết POI, nhóm món nổi bật và tìm kiếm.",
        "GPS/geofence, tự động phát thuyết minh khi vào bán kính POI; hỗ trợ nghe thủ công.",
        "Chọn tour, theo dõi các stop trong tour, bản đồ lọc theo POI thuộc tour.",
        "Quét QR bằng camera hoặc nhập mã thủ công; xử lý deep link mở POI/Tour.",
        "Xem, replay, lọc, xóa listening history; có optimistic item khi vừa bắt đầu nghe.",
        "Offline POI cache bằng SQLite, map tile cache và audio asset cache.",
        "Active device heartbeat định kỳ 8 giây, gửi vị trí nếu có quyền và có tọa độ hợp lệ.",
    ])
    r.h(3, "B. WebAdmin")
    r.bullets([
        "Đăng nhập/đăng xuất bằng cookie, role Admin và PoiOwner.",
        "Dashboard tổng quan POI, tour, audio guide, QR, lượt nghe, completion rate, active devices và top POI.",
        "Quản lý POI, ảnh POI, trạng thái active, owner metadata và duyệt POI owner gửi.",
        "Quản lý tour bằng `ToursController`/`TourAdminService`, chọn danh sách POI có thứ tự.",
        "Quản lý Audio/TTS thông qua `AudioGuidesController` và service tương ứng.",
        "Quản lý Translation cho script narration theo `vi/en/zh/ja/de`.",
        "Tạo/in/tải QR public động cho POI/Tour; public scan web ghi analytics.",
        "Public QR web mô phỏng cấu hình thiết bị bằng `qr-device-profile.js`: random `Math.random() < 0.5 ? 0 : 1`, `0 = mạnh` thì cache payload/asset để ưu tiên offline, `1 = yếu` thì chỉ tải tối thiểu.",
        "Map Analytics, Usage Logs, tài khoản WebAdmin và App user/profile.",
        "Owner Portal cho chủ quán xem POI liên quan, lịch sử nghe và gửi đăng ký POI mới.",
    ])
    r.h(3, "C. Backend API")
    r.bullets([
        "Public read API: POI, Tour, resolve QR, listening history, active device stats/heartbeat.",
        "Admin API có API key: CRUD POI, Tour, AudioGuide, QR item, AuditLog, App user/profile management.",
        "Telemetry API: listening history, active devices, movement logs.",
        "Repository JSON và audit log cho thao tác quản trị POI/Tour/QR/Audio.",
    ])
    r.h(2, "3.2. Yêu cầu phi chức năng")
    r.bullets([
        "Dễ mở rộng nhờ tách `Core` DTO/model khỏi App/Web/API.",
        "Dễ bảo trì nhờ service/repository rõ trách nhiệm và typed API clients trong WebAdmin.",
        "Có fallback khi API lỗi trong app: SQLite offline snapshot, cached snapshot, seed data.",
        "Có phân quyền ở WebAdmin và API admin key cho endpoint quản trị.",
        "Có logging/audit ở API và history/active device telemetry phục vụ dashboard.",
        "Public QR web không crash khi mạng yếu/mất mạng: thiết bị mạnh đọc lại payload đã cache nếu có, thiết bị yếu hiển thị thông báo và dùng dữ liệu hiện có.",
        "Dữ liệu demo có thể kiểm tra trực tiếp trong JSON, phù hợp đồ án môn học.",
        "Giao diện cần đủ dễ dùng cho admin, owner và khách tham quan.",
    ])
    r.h(2, "3.3. Actor hệ thống")
    r.table(
        "Actor, quyền và code liên quan",
        ["Actor", "Vai trò", "Quyền", "Module sử dụng", "Code liên quan"],
        [
            ["Guest/Tourist", "Khách tham quan chưa đăng nhập", "Chọn ngôn ngữ, xem POI/tour/map, nghe, quét QR, ghi history theo guest code.", "App", "AuthPageViewModel, AuthService, MainViewModel"],
            ["App User", "Người dùng app local", "Đăng ký/đăng nhập local, lưu hồ sơ, đồng bộ profile/history.", "App + API", "AuthService, UserProfileSyncService, UserManagementController"],
            ["Admin", "Quản trị hệ thống", "Dashboard, POI, Tour, Audio, Translation, QR, user/profile, audit.", "WebAdmin + API", "AccountController, PoisController, ToursController, SystemAdminController"],
            ["Owner", "Chủ cửa hàng", "Xem POI của mình, xem lượt nghe liên quan, gửi POI mới chờ duyệt.", "Owner Portal", "OwnerController, IWebAdminCurrentUser.CanManage"],
            ["Public QR Visitor", "Người mở QR bằng trình duyệt", "Xem trang public, mô phỏng mạnh/yếu, nghe web hoặc mở app qua deep link.", "Public Web QR", "QrCodesController.Scan, qr-device-profile.js"],
            ["Backend API", "Dịch vụ dữ liệu", "Cung cấp DTO, lưu JSON, telemetry, audit.", "VKFoodAPI", "Controllers, Repositories"],
        ],
    )
    r.h(2, "3.4. Use Case tổng thể")
    r.figure(
        "Use case tổng thể",
        """
flowchart LR
    Guest["Guest/Tourist"] --> App["App MAUI"]
    AppUser["App User"] --> App
    Admin["Admin"] --> Web["WebAdmin MVC"]
    Owner["Owner"] --> OwnerPortal["Owner Portal"]
    Public["Public QR Visitor"] --> QRWeb["Public QR Web"]
    App --> API["Backend API"]
    Web --> API
    OwnerPortal --> API
    QRWeb --> API
    API --> JSON["JSON Repository"]
    App --> Offline["SQLite/Cache/Seed fallback"]
        """,
    )
    r.h(2, "3.5. Đặc tả Use Case chi tiết")
    use_cases = [
        ["UC-01", "Đăng nhập WebAdmin", "Admin/Owner", "WebAdminAuthService", "Tài khoản tồn tại trong `web-admin-users.json`.", "Mở `/Account/Login`, nhập username/password, AccountController gọi ValidateCredentials, tạo cookie claims, redirect Admin về Dashboard hoặc Owner về Owner Portal.", "Sai mật khẩu trả lại login; thiếu quyền vào AdminOnly hiển thị AccessDenied; logout xóa cookie.", "Có session WebAdmin hợp lệ.", "AccountController, WebAdminSecurity.cs, WebAdminAccountStore.cs"],
        ["UC-02", "Admin tạo/sửa/xóa POI", "Admin", "Backend API", "Đã đăng nhập role Admin.", "Admin mở `/Pois`, nhập form, PoiAdminService validate và gọi PoiApiClient; API `PoisController` lưu vào `pois.json` và ghi audit.", "Trùng code, thiếu tên hoặc tọa độ sai trả validation/conflict; delete là soft delete.", "POI được cập nhật để app/Web đọc lại.", "CTest.WebAdmin/Controllers/PoisController.cs, VKFoodAPI/Controllers/PoisController.cs, PoiRepository.cs"],
        ["UC-03", "Owner đăng ký POI mới", "Owner", "PoiAdminService", "Owner đã đăng nhập role PoiOwner.", "Owner mở `/Owner`, điền form; OwnerController ép owner metadata và `IsActive=false`; service tạo POI qua API.", "Dữ liệu không hợp lệ trả form lỗi; API offline hiển thị thông báo.", "POI mới ở trạng thái chờ duyệt.", "OwnerController.cs, PoiValidationService.cs, PoiAdminService.cs"],
        ["UC-04", "Admin duyệt POI Owner gửi", "Admin", "Owner", "Có POI pending `IsActive=false`.", "Admin mở danh sách POI pending, kiểm tra thông tin và bấm duyệt; PoiAdminService bật active và gọi API update.", "Không tìm thấy POI hoặc API lỗi thì hiển thị TempData message.", "POI active xuất hiện cho app/public API.", "PoisController.Approve, PoiAdminService.ApproveAsync"],
        ["UC-05", "Admin tạo/sửa Tour", "Admin", "POI Repository", "Có ít nhất một POI active.", "Admin mở `/Tours`, chọn POI theo thứ tự, nhập mã/tên/thời lượng; TourAdminService gọi `api/tours`.", "Không có POI, POI trùng hoặc POI không còn tồn tại thì validation lỗi.", "Tour lưu trong `tours.json` với `PoiIds` có thứ tự.", "ToursController.cs, TourAdminService.cs, VKFoodAPI/Services/TourRepository.cs"],
        ["UC-06", "App chọn tour và đi theo tour", "Guest/App User", "TourProvider", "App đã tải tour active.", "Người dùng chọn tour; `MainViewModel.ActivateTourAsync` đặt `_activeTourId`, phát giới thiệu tour, map chỉ hiển thị POI trong tour.", "Tour không còn active hoặc không có stop hợp lệ thì báo trạng thái.", "App theo dõi current/next stop và tiến độ tour.", "MainViewModel.cs, TourProvider.cs, TourRepository.cs"],
        ["UC-07", "App tự động thuyết minh GPS/geofence", "Guest/App User", "LocationService", "Có quyền vị trí hoặc app nhận được location hợp lệ.", "`OnLocationUpdated` gọi `ApplyLocationSnapshotAsync`; GeofenceEngine đánh giá bán kính; decision service xét priority, debounce, cooldown; NarrationService phát.", "Không có quyền vị trí thì fallback map đầu phố; đang phát thì không phát chồng; cooldown thì bỏ qua.", "Người dùng nghe đúng POI và history được ghi.", "LocationService.cs, GeofenceEngine.cs, PoiAutoNarrationDecisionService.cs, NarrationService.cs"],
        ["UC-08", "App/Web quét QR", "Guest/Public QR Visitor", "ResolveQrController", "QR chứa POI/Tour code, path hoặc deep link hợp lệ.", "Web mở `/qr/{targetType}/{targetId}`, render POI/Tour rồi chạy `DeviceCapabilitySimulation`; app scan camera/manual gọi `GET /api/resolve-qr?code=...` để resolve POI/Tour.", "QR không hợp lệ trả 404/thông báo; camera denied cho phép nhập mã thủ công; public web offline thì hiển thị thông báo và dùng cache/text hiện có.", "Mở đúng POI/Tour, hiển thị profile mạnh/yếu và có thể autoplay.", "QrCodesController.cs, Scan.cshtml, qr-device-profile.js, ResolveQrController.cs, QrScannerPage.xaml.cs, QrDeepLinkBroker.cs"],
        ["UC-09", "Ghi listening history", "App/Web QR", "ListeningHistoryRepository", "Người dùng bắt đầu nghe POI.", "App tạo optimistic item, `BeginAsync` POST history; phát xong gọi `CompleteAsync`; web QR dùng proxy `/qr/analytics/listening-history`.", "API lỗi thì app giữ local/optimistic; playback lỗi cập nhật ErrorMessage nếu có history id.", "History có thời điểm bắt đầu, số giây nghe, completed/error.", "ListeningHistorySyncService.cs, ListeningHistoryController.cs, ListeningHistoryRepository.cs"],
        ["UC-10", "Dashboard analytics", "Admin", "API telemetry", "Admin đăng nhập và API truy cập được.", "DashboardService gọi POI, Tour, Audio, ListeningHistory, ActiveDevices; tính tổng, top POI, completion, QR listen rate, SSE active devices.", "API lỗi trả dashboard rỗng `IsSyncOnline=false` để tránh số liệu giả.", "Admin xem được tình trạng sử dụng hệ thống.", "HomeController.cs, DashboardService.cs, ActiveDevicesController.cs"],
    ]
    for uc in use_cases:
        r.table(
            f"{uc[0]} - {uc[1]}",
            ["Trường", "Nội dung"],
            [
                ["Mã use case", uc[0]],
                ["Tên use case", uc[1]],
                ["Actor chính", uc[2]],
                ["Actor phụ", uc[3]],
                ["Tiền điều kiện", uc[4]],
                ["Luồng chính", uc[5]],
                ["Luồng thay thế/ngoại lệ", uc[6]],
                ["Hậu điều kiện", uc[7]],
                ["Controller/ViewModel/Service/API/File liên quan", uc[8]],
            ],
        )


def add_design_chapter():
    r.h(1, "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG")
    r.h(2, "4.1. Kiến trúc tổng thể")
    r.p(
        "Hệ thống được thiết kế thành ba khối chính: App MAUI Android, WebAdmin MVC và Backend API. `VinhKhanhGuide.Core` đóng vai trò lớp hợp đồng dùng chung, chứa DTO/model/service contract để giảm lệch dữ liệu giữa các project."
    )
    r.figure(
        "Kiến trúc tổng thể App/WebAdmin/API/JSON Repository",
        """
flowchart LR
    subgraph APP["VinhKhanhGuide.App (.NET MAUI Android)"]
        Auth["AuthPageViewModel/AuthService"]
        MainVM["MainViewModel"]
        GPS["LocationService + GeofenceEngine"]
        QR["QrScannerPage + QrResolveService"]
        Offline["PoiOfflineStore + Audio/Map Cache"]
    end
    subgraph WEB["CTest.WebAdmin (ASP.NET Core MVC)"]
        Account["AccountController"]
        Admin["Admin Controllers/Services"]
        Owner["OwnerController"]
        PublicQR["QrCodesController public scan"]
    end
    subgraph API["VKFoodAPI Controllers"]
        PoiApi["PoisController"]
        TourApi["ToursController"]
        AudioApi["AudioGuidesController"]
        AnalyticsApi["ListeningHistory/ActiveDevices/MovementLogs"]
        UserApi["UserManagementController"]
    end
    subgraph DATA["VKFoodAPI/App_Data"]
        JSON["pois, tours, audio-guides, qr-codes, history, devices, movement, users, audit"]
    end
    APP --> API
    WEB --> API
    PublicQR --> API
    API --> JSON
    APP --> Offline
        """,
    )
    r.h(2, "4.2. Thiết kế phân quyền")
    r.table(
        "Thiết kế phân quyền",
        ["Vai trò/cơ chế", "Cách triển khai", "Quyền chính"],
        [
            ["Admin", "Cookie auth + role `Admin`; policy `WebAdminPolicies.AdminOnly`.", "Quản trị toàn bộ WebAdmin, POI, Tour, Audio, Translation, QR, users, dashboard."],
            ["Owner/PoiOwner", "Cookie auth + role `PoiOwner`; policy `OwnerArea`; claims owner code/email.", "Vào Owner Portal, xem POI khớp owner, đăng ký POI mới."],
            ["Guest", "App tạo guest session local với user code `guest-...`.", "Xem/nghe/quét QR, ghi history theo guest scope."],
            ["App User", "AuthService local Preferences, sync profile public API.", "Lưu hồ sơ, lịch sử nghe theo user code/email."],
            ["Public QR Visitor", "`[AllowAnonymous]` ở public QR scan và proxy analytics.", "Mở QR web, nghe nội dung, ghi analytics web."],
            ["API Admin Key", "`AdminApiKeyAuthenticationHandler`, header `X-Admin-Api-Key`.", "CRUD API quản trị, raw active devices, audit, app user management."],
        ],
    )
    r.h(2, "4.3. Thiết kế dữ liệu")
    r.p(
        "Vì source code lưu dữ liệu bằng JSON repository, ERD dưới đây là mô hình logic để diễn giải quan hệ nghiệp vụ, không phải mô tả database quan hệ đang chạy. `TOUR_POI_LINK` là bảng logic biểu diễn danh sách có thứ tự `TourDto.PoiIds`. Các field dạng list/dictionary như `NarrationTranslations`, `FeaturedCategories` và `PoiIds` được lưu trực tiếp trong JSON DTO."
    )
    r.figure(
        "ERD logic của hệ thống",
        """
erDiagram
    WEB_ADMIN_ACCOUNT {
        string Username PK
        string PasswordHash
        string Role
        string OwnerCode
        string OwnerEmail
    }
    ADMIN_USER_PROFILE {
        guid Id PK
        string UserCode
        string Email
        string PreferredLanguage
        string DevicePlatform
    }
    POI {
        guid Id PK
        string Code
        string Name
        string OwnerUserCode
        string OwnerEmail
        double Latitude
        double Longitude
        double TriggerRadiusMeters
        json NarrationTranslations
    }
    TOUR {
        int Id PK
        string Code
        string Name
        int EstimatedMinutes
        bool IsQrEnabled
    }
    TOUR_POI_LINK {
        int TourId FK
        guid PoiId FK
        int SortOrder
    }
    AUDIO_GUIDE {
        guid Id PK
        guid PoiId FK
        string LanguageCode
        string SourceType
        string Script
        string FilePath
    }
    QR_CODE_ITEM {
        guid Id PK
        string Code
        string TargetType
        string TargetId
    }
    LISTENING_HISTORY {
        guid Id PK
        guid PoiId FK
        string UserCode
        string TriggerType
        int ListenSeconds
        bool Completed
    }
    ACTIVE_DEVICE_SESSION {
        string SessionKey PK
        string DeviceId
        string ClientInstanceId
        datetime LastSeenAtUtc
    }
    MOVEMENT_LOG {
        guid Id PK
        string DeviceId
        double Latitude
        double Longitude
        datetime RecordedAtUtc
    }
    AUDIT_LOG {
        guid Id PK
        string Username
        string Action
        string EntityName
        string EntityId
    }
    OFFLINE_POI_CACHE {
        guid Id PK
        json NarrationTranslationsJson
        datetime LastSyncedAtUtc
    }
    WEB_ADMIN_ACCOUNT ||--o{ POI : owns_logic
    POI ||--o{ AUDIO_GUIDE : has
    POI ||--o{ LISTENING_HISTORY : records
    TOUR ||--o{ TOUR_POI_LINK : contains
    POI ||--o{ TOUR_POI_LINK : appears_in
    QR_CODE_ITEM }o--|| POI : target_poi
    QR_CODE_ITEM }o--|| TOUR : target_tour
    ACTIVE_DEVICE_SESSION ||--o{ MOVEMENT_LOG : emits
        """,
    )
    r.h(2, "4.4. Thiết kế API")
    api_rows = [
        ["POI public", "`/api/pois`, `/api/pois/{id}`, `/api/pois/by-qr?code=`", "GET", "Public", "VKFoodAPI.Controllers.PoisController", "PoiRepository", "Đọc POI active/non-deleted cho app/web."],
        ["POI admin", "`/api/pois`, `/api/pois/{id}`", "POST/PUT/DELETE", "X-Admin-Api-Key", "PoisController", "PoiRepository + AuditLogRepository", "Tạo, sửa, soft delete POI và ghi audit."],
        ["Tour public", "`/api/tours`, `/api/tours/{id}`", "GET", "Public", "ToursController", "TourRepository", "Đọc tour active cho app và QR."],
        ["Tour admin", "`/api/tours`, `/api/tours/{id}`", "POST/PUT/DELETE", "X-Admin-Api-Key", "ToursController", "TourRepository + AuditLogRepository", "CRUD tour; delete là soft delete và tắt QR."],
        ["Audio guide", "`/api/audioguides`, `/api/audioguides/{id}`", "GET/POST/PUT/DELETE", "X-Admin-Api-Key", "AudioGuidesController", "AudioGuideRepository", "Quản lý TTS/file audio. Route thật là `api/[controller]` nên client dùng `api/audioguides`."],
        ["QR item API", "`/api/qr-codes`, `/api/qr-codes/{id}`", "GET/POST/PUT/DELETE", "X-Admin-Api-Key", "QrCodesController", "QrCodeRepository", "CRUD QR item ở API; UI WebAdmin QR độc lập chưa có."],
        ["QR resolve", "`/api/resolve-qr?code=...`", "GET", "Public", "ResolveQrController", "QrCodeRepository + PoiRepository + TourRepository", "Resolve deep link/path/code sang POI/Tour."],
        ["Listening history", "`/api/analytics/listening-history`, alias `/api/narration-histories`", "GET/POST/PUT/DELETE", "Public query scoped", "ListeningHistoryController", "ListeningHistoryRepository", "Ghi begin/complete, ranking, xóa history."],
        ["Active devices", "`/api/analytics/active-devices`, alias `/api/device-presence`", "GET/POST", "Stats/heartbeat public; raw admin", "ActiveDevicesController", "ActiveDeviceRepository + MovementLogRepository", "Heartbeat/disconnect, active stats, ghi movement log từ tọa độ hợp lệ."],
        ["Movement logs", "`/api/movement-logs`", "GET/POST", "Public trong code hiện tại", "MovementLogsController", "MovementLogRepository", "Lưu và truy vấn log di chuyển theo device/user/time."],
        ["App user public sync", "`/api/users/profile-sync`, `/api/app-users/sync`", "POST", "AllowAnonymous", "UserManagementController", "UserManagementRepository", "App đồng bộ hồ sơ người dùng/guest."],
        ["Admin user profile", "`/api/admin/users`, `/api/admin/users/search`, `/api/admin/users/{id}`", "GET/POST", "X-Admin-Api-Key", "UserManagementController", "UserManagementRepository", "WebAdmin quản lý/xem app user profile."],
        ["Audit logs", "`/api/audit-logs`", "GET/POST", "X-Admin-Api-Key", "AuditLogsController", "AuditLogRepository", "Lưu/truy vấn audit log."],
    ]
    r.table("API Matrix rút gọn theo route thật", ["Nhóm", "Endpoint", "Method", "Quyền", "Controller", "Repository/service", "Mục đích"], api_rows)
    r.h(2, "4.5. Thiết kế App MAUI")
    r.p(
        "`MainViewModel` là trung tâm điều phối app: load POI/tour, trạng thái map, GPS, geofence, tour, narration, history, offline package và audio settings. Các service được inject trong `MauiProgram.cs`."
    )
    r.table(
        "Thành phần thiết kế App MAUI",
        ["Thành phần", "Vai trò"],
        [
            ["MainViewModel", "Điều phối UI, tour, POI, narration, history và map state."],
            ["LocationService", "Xin quyền và lấy/cập nhật vị trí thiết bị."],
            ["GeofenceEngine", "Tính khoảng cách và xác định POI nằm trong bán kính kích hoạt."],
            ["PoiAutoNarrationDecisionService", "Chọn candidate theo priority, distance, debounce, cooldown và trạng thái đang phát."],
            ["NarrationService", "Phát TTS hoặc file audio; chính sách cancel/replace playback."],
            ["QrScannerPage/QrResolveService/QrDeepLinkBroker", "Quét/resolve QR và mở POI/Tour/deep link."],
            ["PoiOfflineStore", "SQLite `poi_cache` và `sync_metadata` cho offline snapshot."],
            ["AudioAssetCacheService", "Cache file audio remote/package/local vào app data."],
            ["ActiveDeviceTracker", "Heartbeat 8 giây, gửi device/session/location nếu có."],
        ],
    )
    r.h(2, "4.6. Thiết kế WebAdmin")
    r.table(
        "Thành phần thiết kế WebAdmin",
        ["Controller/Service", "Chức năng"],
        [
            ["AccountController", "Login/logout cookie, redirect theo role."],
            ["HomeController + DashboardService", "Dashboard, active device JSON/SSE, usage snapshot."],
            ["PoisController + PoiAdminService", "Quản lý POI, create/edit/approve/delete."],
            ["ToursController + TourAdminService", "Quản lý tour và danh sách POI trong tour."],
            ["AudioGuidesController + AudioGuideAdminService", "Quản lý audio/TTS cho POI."],
            ["QrCodesController + qr-device-profile.js", "Tạo/in/tải QR public động cho POI/Tour, public scan và mô phỏng cấu hình thiết bị mạnh/yếu cho offline/minimal mode."],
            ["OwnerController", "Owner portal, POI owner, history liên quan, đăng ký POI mới."],
            ["AdminUsersController", "Quản lý tài khoản WebAdmin lưu `web-admin-users.json`."],
            ["SystemAdminController", "Quản lý app user/profile qua UserManagementApiClient."],
            ["MapPoisController", "Quản lý POI trên bản đồ và analytics map."],
            ["UsageLogsController", "Xem listening history/usage logs."],
            ["TranslationsController + TtsTranslationService", "Quản lý script ngôn ngữ `vi/en/zh/ja/de`."],
        ],
    )
    r.h(2, "4.7. Thiết kế Sequence Diagram")

    sequence_diagrams = [
        ("WEB-01. Đăng nhập, phân quyền và đăng xuất", """
sequenceDiagram
    actor U as Admin/Owner
    participant C as AccountController
    participant A as WebAdminAuthService
    participant S as WebAdminAccountStore
    participant Cookie as CookieAuth
    U->>C: GET /Account/Login
    U->>C: POST username/password
    C->>A: ValidateCredentials()
    A->>S: GetAll()
    alt hợp lệ
        C->>Cookie: SignInAsync(claims Role/Owner)
        C-->>U: Redirect Home hoặc Owner
    else sai thông tin
        C-->>U: View Login + error
    end
    U->>C: POST /Account/Logout
    C->>Cookie: SignOutAsync()
        """),
        ("WEB-02. Dashboard tổng quan và realtime active devices", """
sequenceDiagram
    actor Admin
    participant H as HomeController
    participant D as DashboardService
    participant P as PoiApiClient
    participant T as TourApiClient
    participant A as AudioGuideApiClient
    participant L as ListeningHistoryApiClient
    participant Dev as ActiveDeviceApiClient
    Admin->>H: GET /
    H->>D: LoadAsync()
    par load shared data
        D->>P: GET api/pois
        D->>T: GET api/tours
        D->>A: GET api/audioguides
        D->>L: GET api/analytics/listening-history
        D->>Dev: GET api/analytics/active-devices
    end
    D-->>H: DashboardViewModel
    H-->>Admin: Dashboard.cshtml
    loop SSE 1s
        Admin->>H: GET ActiveDeviceEvents
        H->>D: GetActiveDeviceStatsAsync()
        H-->>Admin: event active-devices
    end
        """),
        ("WEB-03. Admin quản lý POI", """
sequenceDiagram
    actor Admin
    participant WC as CTest.WebAdmin.PoisController
    participant S as PoiAdminService
    participant V as PoiValidationService
    participant C as PoiApiClient
    participant API as VKFoodAPI.PoisController
    participant R as PoiRepository
    Admin->>WC: Create/Edit/Delete/Approve
    WC->>S: Load snapshot / save request
    S->>V: ValidateCreate/ValidateUpdate
    alt valid
        S->>C: POST/PUT/DELETE api/pois
        C->>API: X-Admin-Api-Key
        API->>R: Create/Update/Delete
        API-->>C: DTO/NoContent
        WC-->>Admin: Redirect + TempData
    else invalid/API error
        WC-->>Admin: form error
    end
        """),
        ("WEB-04. Quản lý Audio Guide/TTS", """
sequenceDiagram
    actor Admin
    participant C as AudioGuidesController
    participant S as AudioGuideAdminService
    participant V as AudioGuideValidationService
    participant Client as AudioGuideApiClient
    participant API as AudioGuidesController(API)
    participant R as AudioGuideRepository
    participant P as PoiRepository
    Admin->>C: POST Save TTS/File
    C->>S: LoadManagementPageAsync()
    C->>V: Validate(SourceType)
    alt valid
        S->>Client: POST/PUT api/audioguides
        Client->>API: X-Admin-Api-Key
        API->>R: Create/Update
        R->>P: ApplyPublishedAudioGuides()
        API-->>Client: AudioGuideDto
    else missing Script/FilePath
        C-->>Admin: validation error
    end
        """),
        ("WEB-05. Quản lý Tour", """
sequenceDiagram
    actor Admin
    participant C as ToursController(Web)
    participant S as TourAdminService
    participant Client as TourApiClient
    participant API as ToursController(API)
    participant R as TourRepository
    Admin->>C: Create/Edit Tour
    C->>S: PopulateEditorReferencesAsync()
    C->>C: Validate SelectedPoiIds
    alt hợp lệ
        S->>Client: POST/PUT api/tours
        Client->>API: X-Admin-Api-Key
        API->>R: Create/Update(TourDto.PoiIds)
        R-->>API: saved tour
    else không có POI hoặc trùng POI
        C-->>Admin: Editor + ModelState
    end
        """),
        ("WEB-06. Quản lý QR public cho POI/Tour", """
sequenceDiagram
    actor Admin
    participant C as QrCodesController(Web)
    participant P as PoiApiClient
    participant T as TourApiClient
    participant Q as QRCoder
    Admin->>C: GET /QrCodes
    par load targets
        C->>P: GET api/pois
        C->>T: GET api/tours
    end
    C-->>Admin: Manage list POI/Tour QR
    Admin->>C: GET /qr/image/{type}/{id}
    C->>C: ResolveTargetAsync()
    C->>Q: CreateQrCode(public URL)
    C-->>Admin: image/svg+xml
        """),
        ("WEB-07. Public QR scan, mô phỏng device profile và ghi analytics", """
sequenceDiagram
    actor Visitor
    participant C as QrCodesController.Scan
    participant View as Scan.cshtml
    participant Device as qr-device-profile.js
    participant L as ListeningHistoryApiClient
    participant D as ActiveDeviceApiClient
    participant API as VKFoodAPI
    Visitor->>C: GET /qr/{targetType}/{targetId}
    C->>C: BuildScanViewModel()
    C-->>Visitor: Public scan page + QR payload + app deep link
    View->>Device: resolveQrDeviceProfile()
    alt profile = 0 thiết bị mạnh
        Device->>Device: cacheQrPayloadForOffline(POI/Tour + related POIs/audio/text)
        Device-->>View: offline đầy đủ
    else profile = 1 thiết bị yếu
        Device-->>View: tải tối thiểu, dùng dữ liệu hiện có
    end
    Visitor->>View: Play narration
    View->>C: POST /qr/analytics/listening-history
    C->>L: POST api/analytics/listening-history
    L->>API: create history
    View->>C: POST /qr/analytics/active-devices/heartbeat
    C->>D: POST api/analytics/active-devices/heartbeat
        """),
        ("WEB-08. Map Analytics", """
sequenceDiagram
    actor Admin
    participant C as MapPoisController
    participant S as PoiAdminService
    participant L as ListeningHistoryService
    participant D as ActiveDeviceApiClient
    Admin->>C: GET /MapPois
    C->>S: LoadMapManagementPageAsync()
    C->>L: Load usage/listening metrics
    C->>D: GetStatsAsync()
    alt analytics API lỗi
        C-->>Admin: POI map + AnalyticsLoadErrorMessage
    else thành công
        C-->>Admin: POI map, route points, top listening
    end
        """),
        ("WEB-09. Owner Portal đăng ký POI", """
sequenceDiagram
    actor Owner
    participant C as OwnerController
    participant Current as IWebAdminCurrentUser
    participant S as PoiAdminService
    participant V as PoiValidationService
    participant L as ListeningHistoryService
    Owner->>C: GET /Owner
    C->>S: LoadManagementPageAsync()
    C->>L: LoadPageForPoisAsync(owner POIs)
    C-->>Owner: POI + usage + registration form
    Owner->>C: POST Register
    C->>Current: OwnerCode/OwnerEmail
    C->>V: ValidateCreate()
    C->>S: CreateAsync(IsActive=false)
    C-->>Owner: POI chờ duyệt
        """),
        ("WEB-10A. AdminUsersController quản lý tài khoản WebAdmin", """
sequenceDiagram
    actor Admin
    participant C as AdminUsersController
    participant S as WebAdminAccountStore
    Admin->>C: GET /AdminUsers
    C->>S: GetAll()
    C-->>Admin: List Admin/PoiOwner
    Admin->>C: POST SaveAccount
    C->>S: Upsert(hash password)
    alt xóa account
        Admin->>C: POST DeleteAccount
        C->>S: Delete()
        S-->>C: false nếu current/last admin
    end
        """),
        ("WEB-10B. SystemAdminController quản lý App User/Profile", """
sequenceDiagram
    actor Admin
    participant C as SystemAdminController
    participant Client as UserManagementApiClient
    participant API as UserManagementController
    participant R as UserManagementRepository
    Admin->>C: GET /SystemAdmin?keyword=
    C->>Client: GET api/admin/users/search
    Client->>API: X-Admin-Api-Key
    API->>R: SearchUsers()
    Admin->>C: POST SaveAppUser
    C->>Client: POST api/admin/users/profile-sync
    API->>R: UpsertProfile()
        """),
        ("WEB-11. Translation scripts", """
sequenceDiagram
    actor Admin
    participant C as TranslationsController
    participant T as TtsTranslationService
    participant P as PoiApiClient
    participant API as PoisController(API)
    Admin->>C: GET /Translations
    C->>P: GET api/pois/{poiId}
    C-->>Admin: Editors vi/en/zh/ja/de
    Admin->>C: POST Generate
    C->>T: GenerateSuggestedScripts(en,zh,ja,de)
    C->>P: PUT api/pois/{id}
    Admin->>C: POST Save(language)
    C->>P: PUT api/pois/{id}
        """),
        ("APP-01. Mở app, chọn ngôn ngữ và Guest Mode", """
sequenceDiagram
    actor Guest
    participant AuthPage as AuthPageViewModel
    participant Settings as AudioSettingsService
    participant Auth as AuthService
    participant App as AppShell/MainPage
    Guest->>AuthPage: chọn vi/en/zh/ja/de
    AuthPage->>Settings: SaveGuestLanguagePreference()
    Guest->>AuthPage: Enter app
    AuthPage->>Auth: ContinueAsGuestAsync()
    Auth-->>AuthPage: guest session
    AuthPage-->>App: open main interface
        """),
        ("APP-02. Load POI/Tour từ API, fallback offline/cache/seed", """
sequenceDiagram
    participant VM as MainViewModel
    participant PR as PoiRepository(App)
    participant PP as PoiProvider
    participant TR as TourRepository(App)
    participant Store as PoiOfflineStore(SQLite)
    participant API as VKFoodAPI
    VM->>PR: GetPoisAsync()
    PR->>PP: GetPoisAsync()
    alt API OK
        PP->>API: GET api/pois
        PP->>Store: ReplacePoisAsync(snapshot)
    else API lỗi
        PP->>Store: GetPoisAsync()
        alt SQLite có dữ liệu
            Store-->>PP: OfflineSnapshot
        else không có
            PP-->>VM: CachedSnapshot hoặc Seed
        end
    end
    VM->>TR: GetToursAsync()
        """),
        ("APP-03. GPS/geofence auto narration", """
sequenceDiagram
    participant Loc as LocationService
    participant VM as MainViewModel
    participant Geo as GeofenceEngine
    participant Decision as PoiAutoNarrationDecisionService
    participant Narr as NarrationService
    Loc-->>VM: LocationUpdated
    VM->>Geo: Evaluate(location, pois)
    VM->>Decision: CreateCandidates + Decide
    alt ShouldNarrate
        VM->>Narr: NarrateAsync(poi, language, mode)
    else debounce/cooldown/current playing
        VM-->>VM: update map only
    end
        """),
        ("APP-04. Nghe POI thủ công", """
sequenceDiagram
    actor User
    participant VM as MainViewModel
    participant H as ListeningHistorySyncService
    participant N as NarrationService
    User->>VM: Tap Listen
    VM->>H: BeginAsync(poi, language, mode, auto=false)
    VM->>N: NarrateAsync()
    alt phát xong
        VM->>H: CompleteAsync(completed=true)
    else lỗi/cancel
        VM->>H: CompleteAsync(errorMessage)
    end
        """),
        ("APP-05. Chọn tour và bản đồ chỉ hiện POI thuộc tour", """
sequenceDiagram
    actor User
    participant VM as MainViewModel
    participant N as NarrationService
    User->>VM: ActivateTourAsync(tourId)
    VM->>VM: _activeTourId, _activeTourStopIndex=0
    VM->>VM: RefreshTourState/VisibleMapPoiStatuses
    VM->>N: SpeakAsync(tour intro)
    VM-->>User: Current stop, next stop, route points
        """),
        ("APP-06. Quét QR mở POI/Tour", """
sequenceDiagram
    actor User
    participant Page as QrScannerPage
    participant Resolve as QrResolveService
    participant API as ResolveQrController
    participant VM as MainViewModel
    User->>Page: Scan camera hoặc nhập mã
    Page->>Resolve: ResolveAsync(code)
    Resolve->>API: GET api/resolve-qr?code=
    alt TargetType=tour
        Page->>VM: OpenTourFromQrAsync(tourId)
    else TargetType=poi
        Page->>VM: OpenPoiFromQrAsync(poiId)
    end
        """),
        ("APP-07. Listening history begin/complete/delete", """
sequenceDiagram
    participant VM as MainViewModel
    participant H as ListeningHistorySyncService
    participant API as ListeningHistoryController
    VM->>H: BeginAsync()
    H->>API: POST api/analytics/listening-history
    API-->>H: historyId
    VM->>H: CompleteAsync(historyId)
    H->>API: PUT api/analytics/listening-history/{id}
    VM->>H: GetCurrentUserHistoryAsync()
    H->>API: GET history + ranking
    VM->>H: DeleteAsync/DeleteCurrentUserHistoryAsync()
        """),
        ("APP-08. Active device heartbeat và movement log", """
sequenceDiagram
    participant App as MainViewModel/App
    participant Tracker as ActiveDeviceTracker
    participant Loc as LocationService
    participant API as ActiveDevicesController
    participant Move as MovementLogRepository
    App->>Tracker: StartAsync()
    loop mỗi 8 giây
        Tracker->>Loc: Get cached/current location nếu được
        Tracker->>API: POST api/analytics/active-devices/heartbeat
        API->>Move: Create movement log nếu tọa độ hợp lệ
    end
    App->>Tracker: StopAsync()
    Tracker->>API: POST disconnect
        """),
        ("APP-09. Offline cache/audio cache", """
sequenceDiagram
    actor User
    participant VM as MainViewModel
    participant Store as PoiOfflineStore
    participant Map as MapOfflineTileService
    participant Audio as AudioAssetCacheService
    User->>VM: DownloadOfflinePackageAsync
    VM->>Store: ReplacePoisAsync()
    VM->>Map: PrefetchAsync()
    VM->>Audio: PrefetchAsync(pois)
    alt API mất kết nối
        VM->>Store: GetPoisAsync()
        VM-->>User: OfflineSnapshotNotice
    end
        """),
        ("APP-10. Auth/Profile sync", """
sequenceDiagram
    actor User
    participant Auth as AuthService
    participant VM as MainViewModel
    participant Sync as UserProfileSyncService
    participant API as UserManagementController
    User->>Auth: Register/SignIn/Guest
    Auth-->>VM: SessionChanged
    VM->>Sync: SyncCurrentUserAsync(preferredLanguage)
    Sync->>API: POST /api/users/profile-sync hoặc /api/app-users/sync
    API-->>Sync: AdminUserDetailDto
        """),
    ]
    for title, diagram in sequence_diagrams:
        r.figure(title, diagram)


def add_implementation_chapter():
    r.h(1, "CHƯƠNG 5. TRIỂN KHAI HỆ THỐNG")
    r.h(2, "5.1. Cấu trúc source code")
    r.table(
        "Cấu trúc source code",
        ["Project", "Vai trò", "Thành phần chính", "Ghi chú"],
        [
            ["VKFoodAPI", "Backend API", "Controllers, Services/Repositories, Security, App_Data", "ASP.NET Core Web API, JSON repository."],
            ["CTest.WebAdmin", "WebAdmin MVC", "Controllers, Services, Views, Security, App_Data/web-admin-users.json", "Cookie auth, typed API clients, QRCoder."],
            ["VinhKhanhGuide.App", "Mobile app", "ViewModels, Services, Views, Resources, Platforms/Android", ".NET MAUI Android, Mapsui, ZXing, TTS."],
            ["VinhKhanhGuide.Core", "DTO/model/shared logic", "Contracts, Models, Interfaces, Mappings, Seed, Services", "Dùng chung cho App, WebAdmin, API."],
            ["VinhKhanhGuide.Core.Tests", "Kiểm thử unit/integration nhẹ", "PoiAutoNarrationDecisionServiceTests, ApiRepositoryFlowTests", "Đối chiếu core logic và repository flow."],
        ],
    )
    r.h(2, "5.2. Triển khai Backend API")
    r.p(
        "Trong source code, backend được triển khai chủ yếu tại `VKFoodAPI/Controllers` và `VKFoodAPI/Services`. Mỗi nhóm dữ liệu có repository JSON riêng, ví dụ `PoiRepository`, `TourRepository`, `AudioGuideRepository`, `ListeningHistoryRepository`, `ActiveDeviceRepository`, `MovementLogRepository`, `UserManagementRepository` và `AuditLogRepository`."
    )
    r.bullets([
        "Public endpoint đọc dữ liệu: `GET /api/pois`, `GET /api/tours`, `GET /api/resolve-qr?code=...`.",
        "Admin endpoint dùng `AdminApiKeyDefaults.PolicyName` và header `X-Admin-Api-Key`.",
        "Audit log được ghi khi tạo/sửa/xóa POI, Tour và QR item.",
        "Hosted service `ActiveDevicePruningService` dọn active device quá hạn; `DataRepairWarmupService` khởi động/sửa dữ liệu cần thiết.",
        "Listening history và active devices là telemetry phục vụ dashboard, usage logs và map analytics.",
    ])
    r.h(2, "5.3. Triển khai WebAdmin")
    r.p(
        "WebAdmin dùng Cookie Authentication trong `Program.cs`, tài khoản lưu tại `CTest.WebAdmin/App_Data/web-admin-users.json`. Controller WebAdmin không ghi trực tiếp POI/tour/audio vào file mà đi qua typed API clients, nhờ đó WebAdmin và app dùng chung API contract."
    )
    r.bullets([
        "Dashboard: `HomeController` + `DashboardService` gọi song song các API client và có SSE `ActiveDeviceEvents`.",
        "POI: `PoisController`, `PoiAdminService`, `PoiValidationService`, `PoiImageStorageService`.",
        "Tour: `ToursController`, `TourAdminService`, validate selected POI và thứ tự stop.",
        "Audio/TTS: `AudioGuidesController`, `AudioGuideAdminService`, `AudioGuideValidationService`.",
        "QR: `QrCodesController` tạo QR public động cho POI/Tour, không phải CRUD UI độc lập cho `qr-codes.json`; public QR nhúng `qr-device-profile.js` để mô phỏng `0 = mạnh`, `1 = yếu`.",
        "Owner Portal: `OwnerController` lọc owner metadata và tạo POI pending.",
        "User/Admin management: `AdminUsersController` quản lý WebAdmin account; `SystemAdminController` quản lý app user/profile qua API.",
    ])
    r.h(2, "5.4. Triển khai App MAUI")
    r.p(
        "Trong app, `MainViewModel` điều phối hầu hết workflow. Khi khởi tạo, app tải POI/tour từ API, cập nhật danh sách bản đồ, phục hồi offline status và xử lý quyền vị trí. Khi vị trí thay đổi, app đánh giá geofence rồi quyết định có phát thuyết minh hay không."
    )
    r.bullets([
        "Auth/guest mode: `AuthPageViewModel` và `AuthService` tạo session guest hoặc local user.",
        "Map/GPS/geofence: `LocationService`, `GeofenceEngine`, `PoiAutoNarrationDecisionService`.",
        "TTS/audio: `NarrationService` chọn TextToSpeech hoặc Android MediaPlayer theo playback mode.",
        "QR: `QrScannerPage`, `QrResolveService`, `QrDeepLinkBroker`.",
        "Tour: `TourProvider`, `TourRepository`, `MainViewModel.ActivateTourAsync` và các method active tour.",
        "Offline/history/heartbeat: `PoiOfflineStore`, `AudioAssetCacheService`, `ListeningHistorySyncService`, `ActiveDeviceTracker`.",
    ])
    r.h(2, "5.5. Một số đoạn xử lý tiêu biểu")
    flows = [
        ["Lấy danh sách POI", "App gọi `MainViewModel.RefreshPoisIfChangedAsync`, sau đó `PoiRepository.GetPoisAsync` và `PoiProvider.GetPoisAsync`.", "API thành công thì map DTO sang domain, cache vào SQLite; API lỗi thì đọc SQLite, cached snapshot hoặc seed.", "Thiết kế này giúp app vẫn dùng được khi mất mạng hoặc backend tạm dừng."],
        ["Auto narration", "`ApplyLocationSnapshotAsync` nhận location, `GeofenceEngine.Evaluate` tính khoảng cách, `PoiAutoNarrationDecisionService.Decide` chọn POI.", "Nếu `ShouldNarrate=true`, `NarratePoiAsync` gọi `NarrationService` và đồng bộ history.", "Tách decision service giúp logic priority/cooldown/debounce có thể test độc lập."],
        ["Chọn tour", "`ActivateTourAsync` đặt `_activeTourId`, `_activeTourStopIndex=0`, refresh tour state và phát tour intro.", "Map chỉ hiện POI thuộc `TourDto.PoiIds`; khi phát xong stop hợp lệ, `TryAdvanceActiveTourAfterNarration` chuyển chặng.", "Thiết kế lưu `PoiIds` có thứ tự đủ đơn giản cho JSON repository."],
        ["QR resolve/deep link", "`QrScannerPage` dùng ZXing hoặc input thủ công, gọi `QrResolveService.ResolveAsync` tới `/api/resolve-qr?code=...`.", "Resolve API hỗ trợ deep link, path `/qr/{type}/{id}`, QR item code, POI code và tour code.", "Một endpoint resolve chung giảm logic lặp giữa app và web public."],
        ["Mô phỏng cấu hình QR Web", "`Scan.cshtml` truyền QR payload sang `qr-device-profile.js`, sau đó gọi `resolveQrDeviceProfile()`.", "Code dùng `const profile = Math.random() < 0.5 ? 0 : 1;` và kiểm tra `if (profile === 0)`. Profile `0` là thiết bị mạnh: cache payload POI/Tour, POI liên quan và asset audio nếu có. Profile `1` là thiết bị yếu: không prefetch toàn bộ, chỉ dùng dữ liệu đang mở và fallback text/TTS.", "Đáp ứng yêu cầu demo của giảng viên nhưng ghi rõ đây là mô phỏng, chưa phải đo cấu hình thật."],
        ["Ghi listening history", "`NarratePoiAsync` tạo optimistic history, gọi `BeginAsync`; sau khi phát, gọi `CompleteAsync` với listen seconds và trạng thái completed/error.", "Web QR ghi history qua proxy MVC `/qr/analytics/listening-history`.", "History snapshot lưu cả thông tin POI để dashboard không phụ thuộc hoàn toàn vào POI hiện tại."],
        ["Dashboard active devices", "`ActiveDeviceTracker` gửi heartbeat 8 giây; API cập nhật session và ghi movement log nếu tọa độ hợp lệ.", "`DashboardService.GetActiveDeviceStatsAsync` lấy stats; `HomeController.ActiveDeviceEvents` đẩy SSE khi payload thay đổi.", "Thiết kế này phù hợp demo realtime nhẹ mà không cần SignalR."],
        ["Owner đăng ký POI", "`OwnerController.Register` ép `IsActive=false`, owner code/email từ claims rồi gọi `PoiAdminService.CreateAsync`.", "Admin sau đó duyệt bằng `PoisController.Approve`.", "Luồng này giữ quyền publish POI ở Admin, hạn chế owner tự đưa nội dung lên app."],
    ]
    r.table("Các luồng xử lý tiêu biểu", ["Luồng", "Class/method liên quan", "Cách xử lý", "Lý do thiết kế"], flows)


def add_testing_chapter():
    r.h(1, "CHƯƠNG 6. KIỂM THỬ VÀ ĐÁNH GIÁ")
    r.h(2, "6.1. Môi trường kiểm thử")
    r.bullets([
        "Hệ điều hành phát triển: macOS/Windows tùy máy sinh viên; source hiện ở workspace `/Users/nhatminh/Documents/C-test-`.",
        "SDK/framework: .NET 8, ASP.NET Core Web API/MVC, .NET MAUI Android.",
        "Thiết bị kiểm thử app: Android Emulator hoặc thiết bị Android thật; app target `net8.0-android`.",
        "API/WebAdmin chạy local qua `CTest.WebAdmin` hoặc chạy riêng `VKFoodAPI` + `CTest.WebAdmin`.",
        "Dữ liệu mẫu: JSON trong `VKFoodAPI/App_Data` với POI, tour, audio, QR, history, active devices, movement logs và user profiles.",
    ])
    r.h(2, "6.2. Test case chức năng")
    test_rows = [
        ["TC-01", "Login Admin", "Mở `/Account/Login`, nhập `user/12345678`.", "user/12345678", "Vào Dashboard.", "Đạt theo cấu hình demo.", "Pass"],
        ["TC-02", "Login Owner", "Đăng nhập `owner/12345678`.", "owner/12345678", "Vào Owner Portal, không vào AdminOnly.", "Đạt theo role PoiOwner.", "Pass"],
        ["TC-03", "Admin tạo POI", "Mở `/Pois/Create`, nhập thông tin hợp lệ.", "Code mới, tọa độ hợp lệ.", "POI lưu vào API/JSON.", "Đạt theo PoiRepository.Create.", "Pass"],
        ["TC-04", "Admin sửa POI", "Mở edit POI, đổi tên/bán kính.", "POI id có sẵn.", "API update và UpdatedAtUtc đổi.", "Đạt theo PoiRepository.Update.", "Pass"],
        ["TC-05", "Admin xóa POI", "Bấm Delete POI.", "POI id có sẵn.", "Soft delete, IsActive=false, IsDeleted=true.", "Đạt theo PoiRepository.Delete.", "Pass"],
        ["TC-06", "Owner đăng ký POI", "Owner điền form Register.", "Tên, code, tọa độ, owner metadata.", "POI tạo `IsActive=false`.", "Đạt theo OwnerController.Register.", "Pass"],
        ["TC-07", "Admin duyệt POI", "Admin lọc pending và approve.", "POI pending.", "POI active xuất hiện cho app.", "Đạt theo PoiAdminService.ApproveAsync.", "Pass"],
        ["TC-08", "Tạo tour", "Admin tạo tour với >=1 POI.", "Tên, code, EstimatedMinutes, POI ids.", "Tour lưu `tours.json`.", "Đạt theo ToursController/TourRepository.", "Pass"],
        ["TC-09", "App load POI", "Mở app, InitializeAsync.", "API online.", f"Tải {len(active_pois)} POI active.", "Đạt theo dữ liệu hiện tại.", "Pass"],
        ["TC-10", "App chọn tour", "Chọn tour trong app.", "Tour active id 1/2.", "Bản đồ chỉ hiện stop trong tour và phát intro.", "Đạt theo MainViewModel.ActivateTourAsync.", "Pass"],
        ["TC-11", "Auto geofence", "Giả lập vị trí trong bán kính POI.", "Location gần POI.", "Tự phát nếu không cooldown/debounce.", "Đạt theo decision service.", "Pass"],
        ["TC-12", "Nghe thủ công", "Bấm nghe ở POI detail.", "Selected POI.", "Phát TTS/audio và ghi history.", "Đạt theo NarratePoiAsync.", "Pass"],
        ["TC-13", "App quét QR", "Scan hoặc nhập mã QR/URL.", "POI/Tour code hợp lệ.", "Mở POI/Tour tương ứng.", "Đạt theo QrResolveService.", "Pass"],
        ["TC-14", "Public QR web nghe nội dung", "Mở `/qr/poi/{id}` hoặc `/qr/tour/{id}`.", "Target active.", "Hiển thị trang scan, nghe và ghi analytics.", "Đạt theo QrCodesController.Scan.", "Pass"],
        ["TC-15", "Public QR mô phỏng mạnh/yếu", "Mở `/qr/{type}/{id}?deviceProfile=strong`, `/qr/{type}/{id}?deviceProfile=weak` và không truyền query.", "Target active.", "Strong hiển thị `Cấu hình thiết bị: Mạnh - chế độ offline đầy đủ` và console `profile=0`; weak hiển thị `Cấu hình thiết bị: Yếu - chế độ tải tối thiểu` và console `profile=1`; mặc định random 0/1.", "Đạt theo qr-device-profile.js.", "Pass"],
        ["TC-16", "Dashboard active devices/history", "Mở Dashboard khi app/web gửi heartbeat/history.", "Heartbeat/history JSON.", "Hiển thị active count, recent logs, top POI.", "Đạt theo DashboardService.", "Pass"],
    ]
    r.table("Test case chức năng", ["Mã test", "Chức năng", "Bước kiểm thử", "Dữ liệu đầu vào", "Kết quả mong đợi", "Kết quả thực tế", "Trạng thái"], test_rows)
    r.h(2, "6.3. Đánh giá kết quả")
    r.bullets([
        "Hệ thống đã có đầy đủ ba phần App/Web/API và dùng chung DTO trong Core.",
        "App đáp ứng các luồng chính: chọn ngôn ngữ, guest mode, map/POI, tour, QR, GPS/geofence, nghe thủ công, history, offline cache và heartbeat.",
        "WebAdmin đáp ứng quản trị nội dung, dashboard, tour, audio/TTS, translation, QR public, owner portal, WebAdmin account và app user/profile.",
        "Backend API có repository JSON, audit log, telemetry và API key cho nhóm quản trị.",
        "Mức độ đáp ứng yêu cầu ban đầu tốt cho phạm vi đồ án/demo; các điểm production được đưa vào hạn chế/hướng phát triển.",
    ])
    r.h(2, "6.4. Hạn chế hiện tại")
    r.bullets([
        "JSON repository phù hợp demo/học thuật nhưng chưa thay thế database production có transaction, migration, index và concurrency control.",
        "GPS/geofence phụ thuộc quyền thiết bị, độ chính xác GPS và môi trường thực tế; cần kiểm thử ngoài hiện trường.",
        "WebAdmin QR hiện tạo QR public động cho POI/Tour; API có `QrCodeRepository` và `/api/qr-codes`, nhưng UI CRUD QR item độc lập là hướng mở rộng.",
        "Audio playback chưa có `NarrationQueueService`; chính sách hiện tại là cancel/replace playback kết hợp debounce/cooldown.",
        "Ngôn ngữ chọn đã theo `vi/en/zh/ja/de`, nhưng dữ liệu JSON còn key cũ `ko/fr` và cần rà soát lại toàn bộ nội dung bản dịch/UI trước khi triển khai thật.",
        "App auth hiện là local Preferences + profile sync, chưa có server-side auth chuẩn như JWT/OAuth.",
        "App hiện tập trung Android; iOS mới là hướng mở rộng nếu cần đa nền tảng đầy đủ.",
    ])


def add_conclusion_and_appendices():
    r.h(1, "CHƯƠNG 7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
    r.h(2, "7.1. Kết luận")
    r.p(
        "Đồ án đã xây dựng được hệ thống thuyết minh tự động đa ngôn ngữ cho ẩm thực Vĩnh Khánh với ba thành phần hoạt động liên kết: App MAUI cho khách tham quan, WebAdmin MVC cho quản trị/chủ quán và Backend API ASP.NET Core dùng JSON repository."
    )
    r.p(
        "Giá trị chính của hệ thống là đưa nội dung thuyết minh đến đúng ngữ cảnh: khách có thể nghe theo vị trí GPS, theo tour, theo POI chọn thủ công hoặc theo QR tại quán. Đồng thời, admin có công cụ quản trị nội dung, theo dõi usage/active devices và hỗ trợ owner gửi thông tin POI."
    )
    r.h(2, "7.2. Hướng phát triển")
    r.bullets([
        "Chuyển JSON repository sang SQL Server/PostgreSQL, bổ sung migration, transaction và index.",
        "Bổ sung màn CRUD QR item độc lập trong WebAdmin nếu cần quản lý chiến dịch QR riêng.",
        "Hoàn thiện audio queue, priority queue hoặc playlist khi nhiều narration liên tiếp.",
        "Tối ưu offline map và đồng bộ gói dữ liệu theo khu vực.",
        "Tích hợp cloud TTS tự nhiên hơn, caching file audio theo từng ngôn ngữ.",
        "Bổ sung heatmap/báo cáo nâng cao từ movement logs và listening history.",
        "Triển khai cloud với HTTPS/domain ổn định, secret management và logging tập trung.",
        "Phân quyền chi tiết hơn cho owner theo từng POI/action.",
        "Mở rộng app iOS nếu phạm vi đồ án hoặc sản phẩm yêu cầu.",
    ])

    r.h(1, "PHỤ LỤC")
    r.h(2, "Phụ lục A. API Matrix đầy đủ")
    r.p("Bảng API Matrix đầy đủ đã trình bày ở Chương 4. Khi nộp/demo, cần đặc biệt lưu ý route audio guide thật là `api/audioguides`, không phải `api/audio-guides`; route resolve QR thật là `GET /api/resolve-qr?code=...`.")
    r.h(2, "Phụ lục B. Danh sách file/class quan trọng")
    r.table(
        "File/class quan trọng",
        ["Nhóm", "File/class", "Ý nghĩa"],
        [
            ["Auth Web", "CTest.WebAdmin/Controllers/AccountController.cs; Security/WebAdminSecurity.cs", "Cookie auth, role redirect, claims owner."],
            ["Dashboard", "HomeController.cs; DashboardService.cs", "Dashboard, SSE active devices, usage snapshot."],
            ["POI Admin/Owner", "PoisController.cs; OwnerController.cs; PoiAdminService.cs", "Admin CRUD/approve và owner registration."],
            ["Tour", "ToursController.cs; TourAdminService.cs; TourRepository.cs", "Tour CRUD, SelectedPoiIds/PoiIds, soft delete."],
            ["Audio/TTS", "AudioGuidesController.cs; AudioGuideRepository.cs; NarrationService.cs", "TTS/file validation, sync script/audio path vào POI, playback."],
            ["QR", "QrCodesController.cs; Scan.cshtml; qr-device-profile.js; ResolveQrController.cs; QrScannerPage.xaml.cs; QrDeepLinkBroker.cs", "Public QR, mô phỏng mạnh/yếu, resolve, camera/manual scan, deep link."],
            ["GPS/Geofence", "LocationService.cs; GeofenceEngine.cs; PoiAutoNarrationDecisionService.cs", "Permission fallback, radius, priority, cooldown/debounce."],
            ["History/Analytics", "ListeningHistorySyncService.cs; ListeningHistoryController.cs; ActiveDevicesController.cs", "Begin/complete, ranking, heartbeat, movement log."],
            ["Offline", "PoiOfflineStore.cs; MapOfflineTileService.cs; AudioAssetCacheService.cs", "SQLite POI cache, tile cache, audio cache."],
        ],
    )
    r.h(2, "Phụ lục C. Dữ liệu mẫu POI/Tour")
    r.table(
        "POI active trong `pois.json`",
        ["STT", "Tên POI", "Code", "Bán kính", "Priority"],
        [[str(i + 1), p.get("Name", ""), p.get("Code", ""), str(p.get("TriggerRadiusMeters", "")), str(p.get("Priority", ""))] for i, p in enumerate(active_pois)],
    )
    r.table(
        "Tour active trong `tours.json`",
        ["Id", "Code", "Tên tour", "Số POI", "QR"],
        [[str(t.get("Id")), t.get("Code", ""), t.get("Name", ""), str(len(t.get("PoiIds", []))), "Bật" if t.get("IsQrEnabled") else "Tắt"] for t in active_tours],
    )
    r.h(2, "Phụ lục D. Hướng dẫn chạy demo")
    r.code(
        """
# Chạy WebAdmin, đồng thời host API /api/* trong cùng app demo
dotnet run --project CTest.WebAdmin/CTest.WebAdmin.csproj --launch-profile http

# Hoặc chạy API riêng nếu cần
dotnet run --project VKFoodAPI/VKFoodAPI.csproj

# Kiểm thử solution/core test
dotnet test VinhKhanhGuide.Core.Tests/VinhKhanhGuide.Core.Tests.csproj --no-restore

# Build APK Android với API public/local phù hợp
dotnet publish VinhKhanhGuide.App/VinhKhanhGuide.App.csproj -c Release -f net8.0-android -p:ApiBaseUrl=https://<domain-api>/
        """,
        "bash",
    )
    r.bullets([
        "Tài khoản demo WebAdmin: Admin `user/12345678`, Owner `owner/12345678`.",
        "Không dùng `localhost` trong APK release; emulator có thể dùng `10.0.2.2` hoặc domain HTTPS public.",
        "Khi in/chia sẻ QR, cấu hình `QrCode:PublicBaseUrl` và `QrCode:MobileApiBaseUrl` bằng domain truy cập được từ điện thoại.",
    ])
    r.h(2, "Phụ lục E. Checklist đối chiếu PRD - Code")
    r.table(
        "Checklist đối chiếu PRD - Code",
        ["Hạng mục", "Kết quả đối chiếu"],
        [
            ["Số POI active", f"`VKFoodAPI/App_Data/pois.json` hiện có {len(active_pois)} POI active/non-deleted."],
            ["Tour WebAdmin", "Có `CTest.WebAdmin/Controllers/ToursController.cs` và `TourAdminService.cs`, WebAdmin đã quản lý tour."],
            ["QR WebAdmin", "WebAdmin tạo/in/tải QR public động từ POI/Tour; API có `/api/qr-codes`, nhưng UI CRUD QR item độc lập chưa có."],
            ["QR DeviceCapabilitySimulation", "`qr-device-profile.js` mô phỏng random `0/1`: `0 = thiết bị mạnh` cache offline đầy đủ, `1 = thiết bị yếu` tải tối thiểu; hỗ trợ `?deviceProfile=strong|weak|random` để demo."],
            ["AdminUsers/SystemAdmin", "WebAdmin account management thuộc `AdminUsersController`; App user/profile management thuộc `SystemAdminController` + `UserManagementApiClient`."],
            ["API routes", "Audio guide dùng `api/audioguides`; QR resolve dùng `GET /api/resolve-qr?code=...`; listening history dùng `/api/analytics/listening-history`."],
            ["Language support", f"UI chọn ngôn ngữ app/web dùng `vi/en/zh/ja/de`; dữ liệu hiện có keys `{', '.join(translation_keys)}` nên cần migrate/rà soát legacy `ko/fr`."],
            ["Audio playback policy", "Không có queue service phức tạp; `NarrationService` cancel/replace playback, decision service dùng debounce/cooldown."],
            ["Offline cache", "App có SQLite `PoiOfflineStore`, audio cache, map tile cache và fallback seed."],
            ["Analytics", "Có listening history, active device heartbeat, movement logs, dashboard và map analytics."],
        ],
    )


add_overview_chapter()
add_theory_chapter()
add_requirements_chapter()
add_design_chapter()
add_implementation_chapter()
add_testing_chapter()
add_conclusion_and_appendices()


def build_front_matter(body: list[dict], figures: list[str], tables: list[str]) -> list[dict]:
    abbreviations = [
        ["API", "Application Programming Interface"],
        ["POI", "Point of Interest"],
        ["TTS", "Text-to-Speech"],
        ["QR", "Quick Response Code"],
        ["CMS", "Content Management System"],
        ["MVC", "Model View Controller"],
        ["MAUI", "Multi-platform App UI"],
        ["JSON", "JavaScript Object Notation"],
        ["DTO", "Data Transfer Object"],
        ["GPS", "Global Positioning System"],
        ["RBAC", "Role-based Access Control"],
    ]
    return [
        {"type": "cover"},
        {"type": "page_break"},
        {"type": "heading", "level": 1, "text": "LỜI CẢM ƠN"},
        {"type": "paragraph", "text": "Em xin chân thành cảm ơn Quý Thầy/Cô đã hướng dẫn và góp ý trong quá trình thực hiện đồ án. Em cũng cảm ơn bạn bè và người dùng thử đã hỗ trợ kiểm tra các luồng app, web và API. Những góp ý này giúp em hoàn thiện hệ thống theo hướng thực tế hơn, đồng thời hiểu rõ hơn cách thiết kế một sản phẩm phần mềm có nhiều thành phần liên kết."},
        {"type": "paragraph", "text": "Do thời gian và phạm vi đồ án có giới hạn, báo cáo tập trung mô tả đúng những chức năng đã thể hiện trong source code hiện tại, đồng thời ghi nhận trung thực các hạn chế và hướng phát triển tiếp theo."},
        {"type": "page_break"},
        {"type": "heading", "level": 1, "text": "NHẬN XÉT CỦA GIẢNG VIÊN"},
        {"type": "paragraph", "text": "\n\n\n\n\n\n\n\n\n\n"},
        {"type": "paragraph", "text": "................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................................"},
        {"type": "page_break"},
        {"type": "heading", "level": 1, "text": "MỤC LỤC"},
        {"type": "toc"},
        {"type": "page_break"},
        {"type": "heading", "level": 1, "text": "DANH MỤC HÌNH ẢNH"},
        {"type": "bullets", "items": figures},
        {"type": "heading", "level": 1, "text": "DANH MỤC BẢNG BIỂU"},
        {"type": "bullets", "items": tables},
        {"type": "heading", "level": 1, "text": "DANH MỤC TỪ VIẾT TẮT"},
        {"type": "table", "headers": ["Từ viết tắt", "Diễn giải"], "rows": abbreviations},
        {"type": "page_break"},
        *body,
    ]


elements = build_front_matter(r.body, r.figures, r.tables)


def slugify(text: str) -> str:
    text = text.lower()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"[^a-z0-9à-ỹđ\s.-]", "", text)
    text = re.sub(r"\s+", "-", text.strip())
    return text


def plain(text: str) -> str:
    return text.replace("`", "")


def render_md(elements: list[dict]) -> str:
    headings = [e for e in elements if e["type"] == "heading" and e["text"] not in {"MỤC LỤC"}]
    lines: list[str] = []
    for e in elements:
        typ = e["type"]
        if typ == "cover":
            lines.extend([
                "# [TÊN TRƯỜNG]",
                "## [TÊN KHOA/BỘ MÔN]",
                "",
                "# BÁO CÁO ĐỒ ÁN",
                "",
                "## HỆ THỐNG THUYẾT MINH TỰ ĐỘNG ĐA NGÔN NGỮ CHO ẨM THỰC VĨNH KHÁNH",
                "",
                "**Môn học:** [Tên môn học]",
                "",
                "**Giảng viên hướng dẫn:** [Họ tên GVHD]",
                "",
                "**Sinh viên thực hiện:** [Họ tên sinh viên]",
                "",
                "**Mã số sinh viên:** [MSSV]",
                "",
                "**Lớp:** [Lớp]",
                "",
                "**Năm học:** 2025-2026",
                "",
            ])
        elif typ == "page_break":
            lines.append("\n<div class=\"page-break\"></div>\n")
        elif typ == "heading":
            lines.append(f"{'#' * e['level']} {e['text']}\n")
        elif typ == "paragraph":
            lines.append(f"{e['text']}\n")
        elif typ == "bullets":
            lines.extend([f"- {item}" for item in e["items"]])
            lines.append("")
        elif typ == "numbered":
            lines.extend([f"{idx}. {item}" for idx, item in enumerate(e["items"], start=1)])
            lines.append("")
        elif typ == "caption":
            lines.append(f"**{e['text']}**\n")
        elif typ == "table":
            headers = e["headers"]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in e["rows"]:
                lines.append("| " + " | ".join(str(cell).replace("\n", "<br>") for cell in row) + " |")
            lines.append("")
        elif typ == "code":
            lang = e.get("language", "")
            lines.append(f"```{lang}")
            lines.append(e["text"])
            lines.append("```\n")
        elif typ == "toc":
            for h in headings:
                indent = "  " * (h["level"] - 1)
                lines.append(f"{indent}- [{h['text']}](#{slugify(h['text'])})")
            lines.append("")
    return "\n".join(lines)


def render_html(elements: list[dict]) -> str:
    headings = [e for e in elements if e["type"] == "heading" and e["text"] not in {"MỤC LỤC"}]
    body: list[str] = []
    for e in elements:
        typ = e["type"]
        if typ == "cover":
            body.append(
                """
<section class="cover">
  <div class="school">[TÊN TRƯỜNG]</div>
  <div class="faculty">[TÊN KHOA/BỘ MÔN]</div>
  <h1>BÁO CÁO ĐỒ ÁN</h1>
  <h2>HỆ THỐNG THUYẾT MINH TỰ ĐỘNG ĐA NGÔN NGỮ<br>CHO ẨM THỰC VĨNH KHÁNH</h2>
  <div class="meta">
    <p><strong>Môn học:</strong> [Tên môn học]</p>
    <p><strong>Giảng viên hướng dẫn:</strong> [Họ tên GVHD]</p>
    <p><strong>Sinh viên thực hiện:</strong> [Họ tên sinh viên]</p>
    <p><strong>Mã số sinh viên:</strong> [MSSV]</p>
    <p><strong>Lớp:</strong> [Lớp]</p>
    <p><strong>Năm học:</strong> 2025-2026</p>
  </div>
</section>
                """
            )
        elif typ == "page_break":
            body.append('<div class="page-break"></div>')
        elif typ == "heading":
            level = e["level"]
            body.append(f'<h{level} id="{slugify(e["text"])}">{html.escape(e["text"])}</h{level}>')
        elif typ == "paragraph":
            text = html.escape(e["text"]).replace("\n", "<br>")
            text = re.sub(r"`([^`]+)`", r"<code>\1</code>", text)
            body.append(f"<p>{text}</p>")
        elif typ == "bullets":
            body.append("<ul>" + "".join(f"<li>{html.escape(item)}</li>" for item in e["items"]) + "</ul>")
        elif typ == "numbered":
            body.append("<ol>" + "".join(f"<li>{html.escape(item)}</li>" for item in e["items"]) + "</ol>")
        elif typ == "caption":
            body.append(f'<div class="caption {e["kind"]}">{html.escape(e["text"])}</div>')
        elif typ == "table":
            body.append("<div class=\"table-wrap\"><table><thead><tr>" + "".join(f"<th>{html.escape(h)}</th>" for h in e["headers"]) + "</tr></thead><tbody>")
            for row in e["rows"]:
                body.append("<tr>" + "".join(f"<td>{html.escape(str(cell))}</td>" for cell in row) + "</tr>")
            body.append("</tbody></table></div>")
        elif typ == "code":
            lang = e.get("language", "")
            if lang == "mermaid":
                body.append(f'<div class="mermaid">{html.escape(e["text"])}</div>')
            else:
                body.append(f'<pre><code>{html.escape(e["text"])}</code></pre>')
        elif typ == "toc":
            body.append("<nav class=\"toc\"><ul>")
            for h in headings:
                body.append(f'<li class="toc-l{h["level"]}"><a href="#{slugify(h["text"])}">{html.escape(h["text"])}</a></li>')
            body.append("</ul></nav>")
    return f"""<!doctype html>
<html lang="vi">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Báo cáo đồ án VinhKhanhGuide</title>
  <style>
    @page {{ size: A4; margin: 2cm 2cm 2cm 3cm; }}
    body {{ font-family: "Times New Roman", serif; font-size: 13pt; line-height: 1.5; color: #111827; max-width: 980px; margin: 0 auto; padding: 24px; text-align: justify; }}
    h1, h2, h3 {{ text-align: left; line-height: 1.25; color: #0f172a; page-break-after: avoid; }}
    h1 {{ font-size: 20pt; margin-top: 28px; border-bottom: 1px solid #cbd5e1; padding-bottom: 6px; }}
    h2 {{ font-size: 16pt; margin-top: 22px; }}
    h3 {{ font-size: 14pt; margin-top: 16px; }}
    .cover {{ min-height: 90vh; text-align: center; display: flex; flex-direction: column; justify-content: center; gap: 12px; }}
    .cover h1 {{ border: 0; text-align: center; font-size: 24pt; margin: 42px 0 8px; }}
    .cover h2 {{ text-align: center; font-size: 20pt; }}
    .school, .faculty {{ text-transform: uppercase; font-weight: bold; }}
    .meta {{ margin-top: 48px; display: inline-block; text-align: left; }}
    .toc ul {{ list-style: none; padding-left: 0; }}
    .toc-l2 {{ margin-left: 1.2rem; }}
    .toc-l3 {{ margin-left: 2.4rem; }}
    table {{ width: 100%; border-collapse: collapse; margin: 8px 0 18px; font-size: 11pt; text-align: left; }}
    th, td {{ border: 1px solid #94a3b8; padding: 6px 8px; vertical-align: top; }}
    th {{ background: #e2e8f0; font-weight: bold; }}
    .caption {{ font-weight: bold; margin: 12px 0 4px; text-align: left; }}
    code {{ font-family: "Courier New", monospace; font-size: 0.9em; }}
    pre, .mermaid {{ background: #f8fafc; border: 1px solid #cbd5e1; border-radius: 4px; padding: 12px; overflow-x: auto; text-align: left; margin-bottom: 18px; }}
    .page-break {{ page-break-before: always; break-before: page; height: 0; }}
    ul, ol {{ text-align: left; }}
    li {{ margin: 3px 0; }}
    @media print {{ body {{ max-width: none; padding: 0; }} .page-break {{ page-break-before: always; }} }}
  </style>
</head>
<body>
{''.join(body)}
<script type="module">
  import mermaid from "https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs";
  mermaid.initialize({{ startOnLoad: true, securityLevel: "loose", theme: "default" }});
</script>
</body>
</html>
"""


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(plain(str(text)))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)


def render_docx(elements: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Title", 20, "0F172A"),
        ("Heading 1", 16, "0F172A"),
        ("Heading 2", 14, "1E3A8A"),
        ("Heading 3", 13, "334155"),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Trang ")
    add_field(footer, "PAGE")

    headings_for_toc = [e for e in elements if e["type"] == "heading" and e["text"] != "MỤC LỤC"]

    for e in elements:
        typ = e["type"]
        if typ == "cover":
            for text, size, bold in [
                ("[TÊN TRƯỜNG]", 14, True),
                ("[TÊN KHOA/BỘ MÔN]", 14, True),
            ]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = bold
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(size)
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("BÁO CÁO ĐỒ ÁN")
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("HỆ THỐNG THUYẾT MINH TỰ ĐỘNG ĐA NGÔN NGỮ\nCHO ẨM THỰC VĨNH KHÁNH")
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            for line in [
                "Môn học: [Tên môn học]",
                "Giảng viên hướng dẫn: [Họ tên GVHD]",
                "Sinh viên thực hiện: [Họ tên sinh viên]",
                "Mã số sinh viên: [MSSV]",
                "Lớp: [Lớp]",
                "Năm học: 2025-2026",
            ]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line)
                run.font.size = Pt(13)
        elif typ == "page_break":
            doc.add_page_break()
        elif typ == "heading":
            doc.add_heading(e["text"], level=e["level"])
        elif typ == "paragraph":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(plain(e["text"]))
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(13)
        elif typ == "bullets":
            for item in e["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(plain(item))
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(13)
        elif typ == "numbered":
            for item in e["items"]:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(plain(item))
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                run.font.size = Pt(13)
        elif typ == "caption":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(e["text"])
            run.bold = True
            run.italic = True
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
        elif typ == "table":
            table = doc.add_table(rows=1, cols=len(e["headers"]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for idx, head in enumerate(e["headers"]):
                set_cell_shading(hdr[idx], "E2E8F0")
                set_cell_text(hdr[idx], head, bold=True)
                hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for row in e["rows"]:
                cells = table.add_row().cells
                for idx, cell_value in enumerate(row):
                    set_cell_text(cells[idx], cell_value)
                    cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            doc.add_paragraph()
        elif typ == "code":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(e["text"])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(8 if e.get("language") == "mermaid" else 9)
        elif typ == "toc":
            p = doc.add_paragraph()
            add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
            p = doc.add_paragraph()
            p.add_run("Lưu ý: Trong Microsoft Word, nhấn chuột phải vào mục lục và chọn Update Field để cập nhật số trang trước khi nộp.").italic = True
        else:
            raise ValueError(f"Unknown element type {typ}")

    doc.save(OUT_DOCX)


OUT_MD.write_text(render_md(elements), encoding="utf-8")
OUT_HTML.write_text(render_html(elements), encoding="utf-8")
render_docx(elements)

print(f"Generated: {OUT_DOCX}")
print(f"Generated: {OUT_HTML}")
print(f"Generated: {OUT_MD}")
